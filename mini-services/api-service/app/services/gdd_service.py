"""
Gidede — GDD Service
Фаза 4.D.1–4.D.3: Блок 6 — GDD Generator (алгоритм 3.7, Этапы 1–8)

Реализация пайплайна генерации GDD из алгоритма 3.7:
- Этап 1: Определение формата GDD → GDDFormatSpec (3.7.3)
- Этап 2: Маппинг Project State → секции GDD → GDDDataMapping (3.7.4)
- Этап 3: Автозаполнение секций → AutoFilledSections (3.7.5)
- Этап 4: AI-генерация и обогащение → AIEnrichedSections (3.7.6)
- Этап 5: Ручные секции и подсказки → ManualSectionsResult (3.7.7)
- Этап 6: Сшивка и валидация → GDDAssembledDocument (3.7.8)
- Этап 7: Форматирование → GDDFormattedDocument (3.7.9)
- Этап 8: Экспорт → GDDExportResult

Зависимости: 4.A.7 (PromptExecutor), 4.A.8 (Prompt Registry)
"""

import os
import re
import time
import logging
import json
import tempfile
from datetime import datetime
from typing import Any, Optional

from app.ai.executor import PromptExecutor

from app.schemas.gdd import (
    GDDFormat,
    DocAudience,
    DetailLevel,
    GDDFormatSpec,
    SectionMapping,
    SectionReadiness,
    GDDDataMapping,
    SectionContent,
    AutoFilledSections,
    AIEnrichedSections,
    ManualSectionSkeleton,
    ManualSectionsResult,
    GDDGenerationInput,
    GDDProfile,
    SectionPriority,
    ConsistencyIssue,
    ConsistencyReport,
    GDDAssembledSection,
    GDDAssembledDocument,
    GDDFormattedDocument,
    ExportFormat,
    GDDExportResult,
)

logger = logging.getLogger(__name__)


# ============================================================
# Константы: маппинги из алгоритма 3.7
# ============================================================

# Аудитория → формат GDD (алгоритм 3.7.3)
AUDIENCE_FORMAT_MAP: dict[str, str] = {
    "investor": "treatment",
    "team_sync": "sketch_design",
    "production": "full_gdd",
    "personal": "modular",
    "educational": "ten_pager",
}

# Стадия проекта → формат GDD (алгоритм 3.7.3)
STAGE_FORMAT_MAP: dict[str, str] = {
    "concept": "one_sheet",
    "prototype": "ten_pager",
    "preproduction": "sketch_design",
    "production": "full_gdd",
    "live_ops": "modular",
}

# Жанр → уровень детализации (алгоритм 3.7.3)
GENRE_DETAIL_MAP: dict[str, str] = {
    "rpg": "detailed",
    "mmorpg": "exhaustive",
    "strategy": "detailed",
    "tbs": "detailed",
    "rts": "detailed",
    "simulation": "detailed",
    "survival": "detailed",
    "sandbox": "standard",
    "action": "standard",
    "shooter": "standard",
    "fighting": "standard",
    "racing": "standard",
    "metroidvania": "standard",
    "roguelike": "standard",
    "adventure": "standard",
    "horror": "standard",
    "puzzle": "overview",
    "casual": "overview",
    "idle": "overview",
    "party": "overview",
    "educational": "overview",
    "visual_novel": "standard",
    "platformer": "standard",
    "rhythm": "overview",
    "tower_defense": "standard",
}

# Формат → количество страниц (базовое, модифицируется detail_level)
FORMAT_BASE_PAGES: dict[str, int] = {
    "one_sheet": 1,
    "ten_pager": 10,
    "treatment": 3,
    "sketch_design": 15,
    "full_gdd": 50,
    "concept_doc": 8,
    "narrative_bible": 30,
    "modular": 40,
}

# Detail level → множитель страниц
DETAIL_PAGE_MULTIPLIER: dict[str, float] = {
    "overview": 0.5,
    "standard": 1.0,
    "detailed": 1.5,
    "exhaustive": 2.5,
}

# ============================================================
# Шаблоны секций для 8 форматов (алгоритм 3.7.3)
# ============================================================

FORMAT_SECTION_TEMPLATES: dict[str, list[str]] = {
    "one_sheet": [
        "title", "logline", "genre_platform", "target_audience",
        "uniqueness", "visual_hook",
    ],
    "ten_pager": [
        "title_logline", "overview", "core_loop", "mechanics",
        "world_structure", "target_audience", "monetization",
        "milestones_budget", "risks", "team",
    ],
    "treatment": [
        "title", "game_type", "originality", "feasibility",
    ],
    "sketch_design": [
        "mechanics", "level_design", "progression", "hud_ui",
        "content_overview",
    ],
    "full_gdd": [
        # Block 1 (Overview)
        "title", "overview", "genre_platform", "target_audience",
        "uniqueness", "license",
        # Block 2 (Gameplay)
        "core_loop", "controls", "mechanics", "camera_perspective",
        "progression", "balance", "difficulty", "game_modes",
        # Block 3 (Characters/Narrative)
        "characters", "story", "dialogues", "quests", "lore",
        # Block 4 (Levels/World)
        "world_structure", "level_design", "navigation", "combat_spaces",
        # Block 5 (Economy/Progression)
        "resources", "economy", "tech_tree", "difficulty_curve",
        # Block 6 (UI/Visual)
        "hud_ui", "menus", "visual_style", "sound_music",
        # Block 7 (Multiplayer/Social)
        "multiplayer_modes", "social_features", "meta_game",
        # Block 8 (Technical/Business)
        "tech_requirements", "platform_ports", "monetization", "milestones_budget",
    ],
    "concept_doc": [
        "player_experience_goal", "core_loop", "mechanics",
        "system_map", "feedback_patterns", "success_metrics",
    ],
    "narrative_bible": [
        "logline", "story", "lore", "characters", "storyline_map",
        "cutscene_scripts", "dialogues", "quest_matrix", "lore_db",
        "dissension_validator", "pitch_deck",
    ],
    "modular": [
        "concept_overview", "core_loop", "mda_analysis",
        "balance_tables", "progression_curves", "economy_model",
        "world_structure", "character_profiles", "narrative_arc",
        "ui_wireframes", "technical_specs", "monetization", "checklist_results",
    ],
}

# ============================================================
# Маппинг секций → источники данных (алгоритм 3.7.4)
# 38 секций в 8 блоках
# ============================================================

SECTION_DATA_MAP: dict[str, SectionMapping] = {
    # ── Block 1 (Overview) ──────────────────────────────────
    "title": SectionMapping(
        source="concept.title", auto_fill=True, ai_enrich=False,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=False, tables=False, formulas=False,
    ),
    "overview": SectionMapping(
        source="concept.description", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=False, tables=False, formulas=False,
    ),
    "genre_platform": SectionMapping(
        source="concept.genre+concept.platforms", auto_fill=True, ai_enrich=False,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=False, tables=True, formulas=False,
    ),
    "target_audience": SectionMapping(
        source="concept.targetAudience", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=True, manual=False,
        diagram=False, tables=True, formulas=False,
    ),
    "uniqueness": SectionMapping(
        source="concept.usp", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=True, manual=False,
        diagram=False, tables=False, formulas=False,
    ),
    "license": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=False, ai_suggest=True, manual=True,
        diagram=False, tables=False, formulas=False,
    ),
    # ── Block 2 (Gameplay) ──────────────────────────────────
    "core_loop": SectionMapping(
        source="coreLoop", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=True, tables=False, formulas=False,
    ),
    "controls": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=True, tables=False, formulas=False,
    ),
    "mechanics": SectionMapping(
        source="mdaProfile.mechanicSet", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=False, tables=True, formulas=False,
    ),
    "camera_perspective": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=False, formulas=False,
    ),
    "progression": SectionMapping(
        source="progressionProfile", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=True, tables=True, formulas=True,
    ),
    "balance": SectionMapping(
        source="balanceResult", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=True, tables=True, formulas=True,
    ),
    "difficulty": SectionMapping(
        source="progressionProfile.curves.difficulty", auto_fill=True, ai_enrich=False,
        ai_generate=False, ai_suggest=True, manual=False,
        diagram=True, tables=False, formulas=True,
    ),
    "game_modes": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    # ── Block 3 (Characters/Narrative) ──────────────────────
    "characters": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    "story": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=False, formulas=False,
    ),
    "dialogues": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=False, formulas=False,
    ),
    "quests": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    "lore": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=False, formulas=False,
    ),
    # ── Block 4 (Levels/World) ──────────────────────────────
    "world_structure": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=True, tables=False, formulas=False,
    ),
    "level_design": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=True, tables=False, formulas=False,
    ),
    "navigation": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=True, tables=False, formulas=False,
    ),
    "combat_spaces": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=True, tables=False, formulas=False,
    ),
    # ── Block 5 (Economy/Progression) ───────────────────────
    "resources": SectionMapping(
        source="economyProfile.inventory", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=False, tables=True, formulas=False,
    ),
    "economy": SectionMapping(
        source="economyProfile", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=True, tables=True, formulas=True,
    ),
    "tech_tree": SectionMapping(
        source="progressionProfile.contentPlan.unlockTree", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=True, manual=False,
        diagram=True, tables=True, formulas=False,
    ),
    "difficulty_curve": SectionMapping(
        source="progressionProfile.curves.difficulty", auto_fill=True, ai_enrich=False,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=True, tables=False, formulas=True,
    ),
    # ── Block 6 (UI/Visual) ─────────────────────────────────
    "hud_ui": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=True, tables=False, formulas=False,
    ),
    "menus": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    "visual_style": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=False, formulas=False,
    ),
    "sound_music": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=False, formulas=False,
    ),
    # ── Block 7 (Multiplayer/Social) ────────────────────────
    "multiplayer_modes": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    "social_features": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=False, formulas=False,
    ),
    "meta_game": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=False, formulas=False,
    ),
    # ── Block 8 (Technical/Business) ────────────────────────
    "tech_requirements": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    "platform_ports": SectionMapping(
        source="concept.platforms", auto_fill=True, ai_enrich=False,
        ai_generate=False, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    "monetization": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    "milestones_budget": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    # ── Специфичные секции для форматов ─────────────────────
    "logline": SectionMapping(
        source="concept.description", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=True, manual=False,
        diagram=False, tables=False, formulas=False,
    ),
    "visual_hook": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=False, formulas=False,
    ),
    "title_logline": SectionMapping(
        source="concept.title+concept.description", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=False, tables=False, formulas=False,
    ),
    "game_type": SectionMapping(
        source="concept.genre", auto_fill=True, ai_enrich=False,
        ai_generate=False, ai_suggest=True, manual=False,
        diagram=False, tables=False, formulas=False,
    ),
    "originality": SectionMapping(
        source="concept.usp", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=True, manual=False,
        diagram=False, tables=False, formulas=False,
    ),
    "feasibility": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=False, formulas=False,
    ),
    "content_overview": SectionMapping(
        source="progressionProfile.contentPlan", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=True, manual=False,
        diagram=False, tables=True, formulas=False,
    ),
    "player_experience_goal": SectionMapping(
        source="mdaProfile.aesthetics", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=True, manual=False,
        diagram=False, tables=False, formulas=False,
    ),
    "system_map": SectionMapping(
        source="coreLoop+mdaProfile.mechanicSet", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=True, tables=False, formulas=False,
    ),
    "feedback_patterns": SectionMapping(
        source="coreLoop.loops", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=True, manual=False,
        diagram=True, tables=False, formulas=False,
    ),
    "success_metrics": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    "storyline_map": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=True, tables=False, formulas=False,
    ),
    "cutscene_scripts": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=False, formulas=False,
    ),
    "quest_matrix": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    "lore_db": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    "dissension_validator": SectionMapping(
        source="mdaProfile", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=True, manual=False,
        diagram=False, tables=True, formulas=False,
    ),
    "pitch_deck": SectionMapping(
        source="concept", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=True, manual=False,
        diagram=False, tables=True, formulas=False,
    ),
    "risks": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    "team": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=False, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    # ── Modular-специфичные секции ──────────────────────────
    "concept_overview": SectionMapping(
        source="concept", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=False, tables=False, formulas=False,
    ),
    "mda_analysis": SectionMapping(
        source="mdaProfile", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=False, tables=True, formulas=False,
    ),
    "balance_tables": SectionMapping(
        source="balanceResult", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=False, tables=True, formulas=True,
    ),
    "progression_curves": SectionMapping(
        source="progressionProfile.curves", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=True, tables=True, formulas=True,
    ),
    "economy_model": SectionMapping(
        source="economyProfile", auto_fill=True, ai_enrich=True,
        ai_generate=False, ai_suggest=False, manual=False,
        diagram=True, tables=True, formulas=True,
    ),
    "character_profiles": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    "narrative_arc": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=True, tables=False, formulas=False,
    ),
    "ui_wireframes": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=True, tables=False, formulas=False,
    ),
    "technical_specs": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
    "checklist_results": SectionMapping(
        source="", auto_fill=False, ai_enrich=False,
        ai_generate=True, ai_suggest=True, manual=True,
        diagram=False, tables=True, formulas=False,
    ),
}


# ============================================================
# GDD Service
# ============================================================

class GDDService:
    """
    Блок 6: Генератор GDD.
    Реализует алгоритм 3.7 — Этапы 1–5.

    Методы:
    - determine_gdd_format() — Этап 1: определение формата GDD
    - map_project_to_sections() — Этап 2: маппинг Project State → секции
    - generate_auto_sections() — Этап 3: автозаполнение секций
    - generate_ai_sections() — Этап 4: AI-генерация и обогащение секций
    - generate_manual_skeletons() — Этап 5: ручные секции с подсказками
    - generate_stages_1_3() — пайплайн Этапов 1–3
    - generate_stages_1_5() — полный пайплайн Этапов 1–5
    """

    def __init__(self, executor: PromptExecutor):
        self.executor = executor

    # ========================================================
    # Этап 1: Определение формата GDD (3.7.3)
    # ========================================================

    async def determine_gdd_format(
        self,
        input_data: GDDGenerationInput,
    ) -> GDDFormatSpec:
        """
        Этап 1: Определение формата GDD на основе входных данных.

        Алгоритм 3.7.3:
        1. Если targetFormat указан → используем его
        2. Иначе → эвристика по audience и project stage
        3. Определяем detail_level по жанру
        4. Определяем секции по формату
        5. Оцениваем количество страниц

        Returns:
            GDDFormatSpec с форматом, секциями и оценкой страниц
        """
        start = time.time()

        # === Шаг 1.1: Определение формата ===
        target_format = input_data.target_format

        if not target_format:
            # Эвристика: audience имеет приоритет
            if input_data.target_audience_doc:
                target_format = AUDIENCE_FORMAT_MAP.get(
                    input_data.target_audience_doc, "full_gdd"
                )
                logger.info(
                    f"[Stage 1.1] Format from audience "
                    f"'{input_data.target_audience_doc}': {target_format}"
                )
            elif input_data.project_stage:
                target_format = STAGE_FORMAT_MAP.get(
                    input_data.project_stage, "full_gdd"
                )
                logger.info(
                    f"[Stage 1.1] Format from project stage "
                    f"'{input_data.project_stage}': {target_format}"
                )
            else:
                target_format = "full_gdd"
                logger.info("[Stage 1.1] Default format: full_gdd")

        # === Шаг 1.2: Определение уровня детализации ===
        detail_level = input_data.detail_level

        if not detail_level:
            genre = ""
            if input_data.concept:
                genre = input_data.concept.get("genre", "").lower()

            if genre:
                detail_level = GENRE_DETAIL_MAP.get(genre, "standard")
                logger.info(
                    f"[Stage 1.2] Detail level from genre '{genre}': {detail_level}"
                )
            else:
                detail_level = "standard"
                logger.info("[Stage 1.2] Default detail level: standard")

        # === Шаг 1.3: Определение секций ===
        sections = list(FORMAT_SECTION_TEMPLATES.get(target_format, []))

        # Добавляем пользовательские секции
        if input_data.custom_sections:
            for section in input_data.custom_sections:
                if section not in sections:
                    sections.append(section)

        # Удаляем исключённые секции
        if input_data.excluded_sections:
            sections = [
                s for s in sections
                if s not in input_data.excluded_sections
            ]

        # === Шаг 1.4: Оценка страниц ===
        base_pages = FORMAT_BASE_PAGES.get(target_format, 20)
        page_multiplier = DETAIL_PAGE_MULTIPLIER.get(detail_level, 1.0)
        estimated_pages = max(1, int(base_pages * page_multiplier))

        # === Сборка результата ===
        format_spec = GDDFormatSpec(
            format=target_format,
            detail_level=detail_level,
            sections=sections,
            estimated_pages=estimated_pages,
            audience=input_data.target_audience_doc,
        )

        logger.info(
            f"[Stage 1] Format determined: {target_format}, "
            f"detail={detail_level}, sections={len(sections)}, "
            f"pages≈{estimated_pages} ({time.time() - start:.2f}s)"
        )

        return format_spec

    # ========================================================
    # Этап 2: Маппинг Project State → секции GDD (3.7.4)
    # ========================================================

    async def map_project_to_sections(
        self,
        format_spec: GDDFormatSpec,
        input_data: GDDGenerationInput,
    ) -> GDDDataMapping:
        """
        Этап 2: Маппинг Project State → секции GDD.

        Алгоритм 3.7.4:
        1. Фильтруем маппинги по активным секциям из format_spec
        2. Проверяем готовность каждой секции (есть ли данные в input?)
        3. Классифицируем секции: auto_fillable, manual, ai_generatable
        4. Рассчитываем coverage_score

        Returns:
            GDDDataMapping с маппингами, готовностью и покрытием
        """
        start = time.time()

        active_sections = format_spec.sections
        active_mappings: dict[str, SectionMapping] = {}
        section_readiness: dict[str, SectionReadiness] = {}
        auto_fillable_sections: list[str] = []
        manual_sections: list[str] = []
        ai_generatable_sections: list[str] = []

        for section_name in active_sections:
            # Получаем маппинг для секции (или создаём default)
            mapping = SECTION_DATA_MAP.get(
                section_name,
                SectionMapping(
                    source="", auto_fill=False, ai_enrich=False,
                    ai_generate=True, ai_suggest=True, manual=True,
                    diagram=False, tables=False, formulas=False,
                ),
            )
            active_mappings[section_name] = mapping

            # Проверяем готовность секции
            readiness = self._check_section_readiness(
                section_name, mapping, input_data
            )
            section_readiness[section_name] = readiness

            # Классифицируем секцию
            if readiness.status == "ready" and mapping.auto_fill:
                auto_fillable_sections.append(section_name)
            elif readiness.status == "ai_generatable" or mapping.ai_generate:
                ai_generatable_sections.append(section_name)
            elif readiness.status == "ai_suggestable" and mapping.ai_suggest:
                ai_generatable_sections.append(section_name)
            else:
                manual_sections.append(section_name)

        # Рассчитываем coverage_score
        total = len(active_sections)
        auto_count = len(auto_fillable_sections)
        coverage_score = round(auto_count / total, 3) if total > 0 else 0.0

        data_mapping = GDDDataMapping(
            format_spec=format_spec,
            active_mappings=active_mappings,
            section_readiness=section_readiness,
            auto_fillable_sections=auto_fillable_sections,
            manual_sections=manual_sections,
            ai_generatable_sections=ai_generatable_sections,
            coverage_score=coverage_score,
        )

        logger.info(
            f"[Stage 2] Mapping complete: {len(active_mappings)} sections, "
            f"auto_fillable={auto_count}, manual={len(manual_sections)}, "
            f"ai_generatable={len(ai_generatable_sections)}, "
            f"coverage={coverage_score:.2f} ({time.time() - start:.2f}s)"
        )

        return data_mapping

    def _check_section_readiness(
        self,
        section_name: str,
        mapping: SectionMapping,
        input_data: GDDGenerationInput,
    ) -> SectionReadiness:
        """
        Проверка готовности секции на основе наличия данных в input.

        Возвращает:
            SectionReadiness со статусом и уровнем покрытия
        """
        source = mapping.source
        if not source:
            # Нет источника — требуется ручное заполнение или AI
            if mapping.ai_generate:
                return SectionReadiness(
                    status="ai_generatable",
                    coverage=0.0,
                    auto_fillable=False,
                )
            return SectionReadiness(
                status="manual_required",
                coverage=0.0,
                auto_fillable=False,
            )

        # Проверяем наличие данных по источнику
        coverage = self._assess_source_coverage(source, input_data)

        if coverage >= 0.8:
            status = "ready"
            auto_fillable = mapping.auto_fill
        elif coverage >= 0.4:
            status = "ai_suggestable"
            auto_fillable = mapping.auto_fill and coverage >= 0.6
        elif mapping.ai_generate:
            status = "ai_generatable"
            auto_fillable = False
        else:
            status = "manual_required"
            auto_fillable = False

        return SectionReadiness(
            status=status,
            coverage=round(coverage, 3),
            auto_fillable=auto_fillable,
        )

    def _assess_source_coverage(
        self,
        source: str,
        input_data: GDDGenerationInput,
    ) -> float:
        """
        Оценка покрытия данных для источника.

        Проверяет наличие и полноту данных по пути source
        во входных данных (concept, coreLoop, mdaProfile, etc.)
        """
        # Разбиваем source на части (напр. "concept.title" → ["concept", "title"])
        parts = source.split("+")
        total_coverage = 0.0

        for part in parts:
            part = part.strip()
            coverage = self._assess_single_source(part, input_data)
            total_coverage = max(total_coverage, coverage)

        return total_coverage

    def _assess_single_source(
        self,
        source: str,
        input_data: GDDGenerationInput,
    ) -> float:
        """Оценка покрытия для одного источника."""
        path_parts = source.split(".")
        root = path_parts[0]
        sub_path = path_parts[1:] if len(path_parts) > 1 else []

        # Получаем корневой объект
        root_data: Optional[dict] = None
        if root == "concept" and input_data.concept:
            root_data = input_data.concept
        elif root == "coreLoop" and input_data.core_loop:
            root_data = input_data.core_loop
        elif root == "mdaProfile" and input_data.mda_profile:
            root_data = input_data.mda_profile
        elif root == "balanceResult" and input_data.balance_result:
            root_data = input_data.balance_result
        elif root == "progressionProfile" and input_data.progression_profile:
            root_data = input_data.progression_profile
        elif root == "economyProfile" and input_data.economy_profile:
            root_data = input_data.economy_profile

        if root_data is None:
            return 0.0

        # Если нет подпути — корневой объект существует
        if not sub_path:
            # Проверяем что объект не пустой
            if isinstance(root_data, dict) and len(root_data) > 0:
                return 1.0
            return 0.3  # Пустой объект

        # Проходим по подпути
        current = root_data
        for part in sub_path:
            if isinstance(current, dict):
                if part in current:
                    current = current[part]
                else:
                    return 0.2  # Подпуть не найден
            else:
                return 0.1

        # Финальное значение существует
        if current is not None and current != "" and current != []:
            return 1.0
        return 0.3

    # ========================================================
    # Этап 3: Автозаполнение секций (3.7.5)
    # ========================================================

    async def generate_auto_sections(
        self,
        data_mapping: GDDDataMapping,
        input_data: GDDGenerationInput,
    ) -> AutoFilledSections:
        """
        Этап 3: Автозаполнение секций GDD из Project State.

        Алгоритм 3.7.5:
        1. Для каждой auto-fillable секции где данные готовы:
           - Извлечь исходные данные из input
           - Отформатировать как Markdown
        2. Типы форматирования:
           - Текстовые секции → прямое форматирование
           - Секции с диаграммами → текст + markdown-диаграмма
           - Секции с таблицами → markdown-таблица
           - Секции с формулами → форматированные формулы

        Returns:
            AutoFilledSections с заполненными секциями
        """
        start = time.time()

        sections: dict[str, SectionContent] = {}
        auto_fillable = data_mapping.auto_fillable_sections

        for section_name in auto_fillable:
            mapping = data_mapping.active_mappings.get(section_name)
            if not mapping:
                continue

            readiness = data_mapping.section_readiness.get(section_name)
            if not readiness or not readiness.auto_fillable:
                continue

            # Извлекаем и форматируем данные
            content = self._format_section_content(
                section_name, mapping, input_data
            )

            if content:
                sections[section_name] = content

        # Рассчитываем общее покрытие
        total_sections = len(data_mapping.format_spec.sections)
        filled_count = len(sections)
        total_coverage = round(filled_count / total_sections, 3) if total_sections > 0 else 0.0

        result = AutoFilledSections(
            sections=sections,
            count=filled_count,
            total_coverage=total_coverage,
        )

        logger.info(
            f"[Stage 3] Auto-filled: {filled_count}/{total_sections} sections, "
            f"coverage={total_coverage:.2f} ({time.time() - start:.2f}s)"
        )

        return result

    def _format_section_content(
        self,
        section_name: str,
        mapping: SectionMapping,
        input_data: GDDGenerationInput,
    ) -> Optional[SectionContent]:
        """
        Форматирование содержимого секции на основе типа и источника.

        Возвращает SectionContent или None если данных недостаточно
        """
        source = mapping.source
        if not source:
            return None

        # Извлекаем данные из источника
        raw_data = self._extract_source_data(source, input_data)

        if raw_data is None:
            return None

        # Форматируем по типу секции
        formatter = self._get_section_formatter(section_name)
        content_text = formatter(raw_data, section_name, mapping)

        # Обрабатываем дополнительные элементы
        diagram = None
        tables = None
        formulas = None
        requires_review = False

        if mapping.diagram:
            diagram = self._generate_diagram_markdown(section_name, raw_data)

        if mapping.tables:
            tables = self._generate_tables(section_name, raw_data)

        if mapping.formulas:
            formulas = self._extract_formulas(section_name, raw_data)

        # Секции с AI-enrich требуют ревью
        if mapping.ai_enrich:
            requires_review = True

        return SectionContent(
            content=content_text,
            source="auto_fill",
            auto_filled=True,
            diagram=diagram,
            tables=tables,
            formulas=formulas,
            requires_review=requires_review,
        )

    def _extract_source_data(
        self,
        source: str,
        input_data: GDDGenerationInput,
    ) -> Any:
        """Извлечение данных из источника по пути source."""
        # Для составных источников (через +) объединяем
        parts = source.split("+")
        if len(parts) > 1:
            combined = {}
            for part in parts:
                data = self._extract_single_source(part.strip(), input_data)
                if isinstance(data, dict):
                    combined.update(data)
                elif data is not None:
                    combined[part.strip().split(".")[-1]] = data
            return combined if combined else None

        return self._extract_single_source(parts[0].strip(), input_data)

    def _extract_single_source(
        self,
        source: str,
        input_data: GDDGenerationInput,
    ) -> Any:
        """Извлечение данных из одного источника."""
        path_parts = source.split(".")
        root = path_parts[0]
        sub_path = path_parts[1:] if len(path_parts) > 1 else []

        root_data: Optional[dict] = None
        if root == "concept" and input_data.concept:
            root_data = input_data.concept
        elif root == "coreLoop" and input_data.core_loop:
            root_data = input_data.core_loop
        elif root == "mdaProfile" and input_data.mda_profile:
            root_data = input_data.mda_profile
        elif root == "balanceResult" and input_data.balance_result:
            root_data = input_data.balance_result
        elif root == "progressionProfile" and input_data.progression_profile:
            root_data = input_data.progression_profile
        elif root == "economyProfile" and input_data.economy_profile:
            root_data = input_data.economy_profile

        if root_data is None:
            return None

        if not sub_path:
            return root_data

        # Проходим по подпути
        current: Any = root_data
        for part in sub_path:
            if isinstance(current, dict) and part in current:
                current = current[part]
            else:
                return None

        return current

    def _get_section_formatter(self, section_name: str):
        """Получить функцию форматирования для секции."""
        formatters = {
            "title": self._format_text_field,
            "logline": self._format_text_field,
            "overview": self._format_rich_text,
            "genre_platform": self._format_genre_platform,
            "target_audience": self._format_text_field,
            "uniqueness": self._format_text_field,
            "core_loop": self._format_core_loop,
            "mechanics": self._format_mechanics,
            "progression": self._format_progression,
            "balance": self._format_balance,
            "difficulty": self._format_difficulty,
            "resources": self._format_resources,
            "economy": self._format_economy,
            "tech_tree": self._format_tech_tree,
            "difficulty_curve": self._format_difficulty_curve,
            "title_logline": self._format_text_field,
            "game_type": self._format_text_field,
            "originality": self._format_text_field,
            "content_overview": self._format_content_overview,
            "player_experience_goal": self._format_text_field,
            "system_map": self._format_system_map,
            "feedback_patterns": self._format_feedback_patterns,
            "dissension_validator": self._format_dissension,
            "pitch_deck": self._format_pitch_deck,
            "concept_overview": self._format_concept_overview,
            "mda_analysis": self._format_mda_analysis,
            "balance_tables": self._format_balance,
            "progression_curves": self._format_progression,
            "economy_model": self._format_economy,
            "platform_ports": self._format_genre_platform,
            "visual_hook": self._format_text_field,
            "feasibility": self._format_text_field,
            "license": self._format_text_field,
            "success_metrics": self._format_text_field,
        }
        return formatters.get(section_name, self._format_generic)

    # ========================================================
    # Форматтеры секций
    # ========================================================

    def _format_text_field(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование простого текстового поля."""
        if isinstance(data, str):
            return data
        elif isinstance(data, dict):
            # Пытаемся извлечь осмысленный текст
            text_keys = [
                "title", "name", "description", "text", "logline",
                "usp", "value", "content",
            ]
            for key in text_keys:
                if key in data and isinstance(data[key], str):
                    return data[key]
            # Если не нашли — формат как список ключ-значение
            lines = []
            for k, v in data.items():
                if isinstance(v, (str, int, float, bool)):
                    lines.append(f"**{k}**: {v}")
            return "\n".join(lines) if lines else str(data)
        return str(data) if data is not None else ""

    def _format_rich_text(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование расширенного текстового содержимого."""
        if isinstance(data, str):
            return data
        elif isinstance(data, dict):
            lines = []
            # Заголовок если есть
            for title_key in ["title", "name"]:
                if title_key in data and isinstance(data[title_key], str):
                    lines.append(f"## {data[title_key]}")
                    break

            # Описание
            for desc_key in ["description", "overview", "text", "summary"]:
                if desc_key in data and isinstance(data[desc_key], str):
                    lines.append(data[desc_key])
                    break

            # Остальные поля как список
            skip_keys = {"title", "name", "description", "overview", "text", "summary"}
            for k, v in data.items():
                if k in skip_keys:
                    continue
                if isinstance(v, (str, int, float, bool)):
                    lines.append(f"- **{k}**: {v}")
                elif isinstance(v, list):
                    lines.append(f"- **{k}**: {', '.join(str(i) for i in v[:10])}")

            return "\n\n".join(lines) if lines else str(data)
        return str(data) if data is not None else ""

    def _format_genre_platform(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование жанра и платформ."""
        if isinstance(data, dict):
            lines = []
            if "genre" in data:
                lines.append(f"**Жанр**: {data['genre']}")
            if "subgenre" in data:
                lines.append(f"**Поджанр**: {data['subgenre']}")
            if "platforms" in data:
                platforms = data["platforms"]
                if isinstance(platforms, list):
                    lines.append(f"**Платформы**: {', '.join(str(p) for p in platforms)}")
                else:
                    lines.append(f"**Платформы**: {platforms}")
            # Остальные поля
            skip = {"genre", "subgenre", "platforms"}
            for k, v in data.items():
                if k not in skip and isinstance(v, (str, int, float)):
                    lines.append(f"**{k}**: {v}")
            return "\n".join(lines) if lines else str(data)
        elif isinstance(data, list):
            # Платформы как список
            return f"**Платформы**: {', '.join(str(p) for p in data)}"
        return str(data) if data is not None else ""

    def _format_core_loop(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование Core Loop с диаграммой."""
        if not isinstance(data, dict):
            return str(data) if data is not None else ""

        lines = []

        # Структурный тип
        struct_type = data.get("structural_type", "")
        if isinstance(struct_type, dict):
            st_name = struct_type.get("type", struct_type.get("name", ""))
            lines.append(f"**Структурный тип**: {st_name}")
        elif isinstance(struct_type, str) and struct_type:
            lines.append(f"**Структурный тип**: {struct_type}")

        # Шаги
        steps = data.get("steps", [])
        if steps and isinstance(steps, list):
            lines.append("\n### Шаги Core Loop")
            for i, step in enumerate(steps):
                if isinstance(step, dict):
                    name = step.get("name", step.get("action", f"Шаг {i+1}"))
                    lines.append(f"{i+1}. **{name}**")
                    if "resources_consumed" in step:
                        consumed = step["resources_consumed"]
                        if isinstance(consumed, list) and consumed:
                            lines.append(f"   - Потребляет: {', '.join(str(r) for r in consumed[:5])}")
                    if "resources_produced" in step:
                        produced = step["resources_produced"]
                        if isinstance(produced, list) and produced:
                            lines.append(f"   - Производит: {', '.join(str(r) for r in produced[:5])}")
                elif isinstance(step, str):
                    lines.append(f"{i+1}. {step}")

        # Ресурсы
        resources = data.get("resources", [])
        if resources and isinstance(resources, list):
            lines.append("\n### Ресурсы")
            for res in resources[:10]:
                if isinstance(res, dict):
                    name = res.get("name", "")
                    lines.append(f"- **{name}**")
                elif isinstance(res, str):
                    lines.append(f"- {res}")

        return "\n".join(lines) if lines else ""

    def _format_mechanics(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование механик."""
        if isinstance(data, dict):
            mechanic_set = data.get("mechanicSet", data.get("mechanic_set", data))
            if isinstance(mechanic_set, dict):
                mechanics = mechanic_set.get("mechanics", mechanic_set.get("base", []))
            else:
                mechanics = []
        elif isinstance(data, list):
            mechanics = data
        else:
            return str(data) if data is not None else ""

        lines = []
        if isinstance(mechanics, list):
            for mech in mechanics[:15]:
                if isinstance(mech, dict):
                    name = mech.get("name", mech.get("mechanic_name", ""))
                    desc = mech.get("description", "")
                    group = mech.get("group", mech.get("group_name", ""))
                    if name:
                        line = f"- **{name}**"
                        if group:
                            line += f" ({group})"
                        if desc:
                            line += f": {desc[:100]}"
                        lines.append(line)
                elif isinstance(mech, str):
                    lines.append(f"- {mech}")

        return "\n".join(lines) if lines else ""

    def _format_progression(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование прогрессии."""
        if not isinstance(data, dict):
            return str(data) if data is not None else ""

        lines = []

        # Макро-параметры
        macro = data.get("macroModel", data)
        if isinstance(macro, dict):
            duration = macro.get("duration", 0)
            levels = macro.get("totalLevels", macro.get("levels", 0))
            prog_type = macro.get("progressionType", "")
            if duration:
                lines.append(f"**Длительность**: {duration} ч.")
            if levels:
                lines.append(f"**Уровни**: {levels}")
            if prog_type:
                lines.append(f"**Тип прогрессии**: {prog_type}")

        # Тиры
        tier_model = data.get("tierModel", {})
        if isinstance(tier_model, dict):
            tiers = tier_model.get("tiers", [])
            if tiers:
                lines.append("\n### Тиры")
                for tier in tiers:
                    if isinstance(tier, dict):
                        idx = tier.get("index", 0)
                        scale = tier.get("scale", "")
                        levels = tier.get("level_count", 0)
                        lines.append(
                            f"- **Тир {idx+1}** ({scale}): {levels} уровней"
                        )

        # Кривые
        curves = data.get("curves", {})
        if isinstance(curves, dict):
            xp = curves.get("xp_to_level", {})
            if isinstance(xp, dict) and xp.get("formula"):
                lines.append(f"\n**XP-кривая**: {xp['formula']}")

        return "\n".join(lines) if lines else ""

    def _format_balance(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование балансировки."""
        if not isinstance(data, dict):
            return str(data) if data is not None else ""

        lines = []

        # Тип балансировки
        balance_type = data.get("balanceType", data.get("dominant_type", ""))
        if balance_type:
            lines.append(f"**Тип балансировки**: {balance_type}")

        # Транзитивная балансировка
        transitive = data.get("transitiveResult", data.get("transitive", {}))
        if isinstance(transitive, dict) and transitive:
            lines.append("\n### Транзитивная балансировка")
            score = transitive.get("correlation", transitive.get("score", 0))
            lines.append(f"- Корреляция cost/power: {score:.3f}")

        # Нетранзитивная балансировка
        intransitive = data.get("intransitiveResult", data.get("intransitive", {}))
        if isinstance(intransitive, dict) and intransitive:
            lines.append("\n### Нетранзитивная балансировка")
            rps = intransitive.get("rpsCycles", intransitive.get("cycles", []))
            if isinstance(rps, list) and rps:
                lines.append(f"- RPS-циклов: {len(rps)}")

        return "\n".join(lines) if lines else ""

    def _format_difficulty(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование кривой сложности."""
        if isinstance(data, dict):
            formula = data.get("formula", "")
            params = data.get("parameters", {})
            lines = []
            if formula:
                lines.append(f"**Формула**: `{formula}`")
            if isinstance(params, dict) and params:
                lines.append("**Параметры**:")
                for k, v in params.items():
                    if isinstance(v, (int, float, str)):
                        lines.append(f"- {k}: {v}")
            return "\n".join(lines) if lines else ""
        return str(data) if data is not None else ""

    def _format_resources(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование ресурсов экономики."""
        if isinstance(data, dict):
            inventory = data.get("inventory", data)
            if isinstance(inventory, dict):
                resources = inventory.get("resources", [])
            else:
                resources = []
        elif isinstance(data, list):
            resources = data
        else:
            return str(data) if data is not None else ""

        lines = []
        if isinstance(resources, list):
            for res in resources[:15]:
                if isinstance(res, dict):
                    name = res.get("name", "")
                    res_class = res.get("resource_class", "")
                    consumable = res.get("is_consumable", False)
                    lines.append(
                        f"- **{name}** ({res_class})"
                        + (" — потребляемый" if consumable else "")
                    )
                elif isinstance(res, str):
                    lines.append(f"- {res}")

        return "\n".join(lines) if lines else ""

    def _format_economy(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование экономики."""
        if not isinstance(data, dict):
            return str(data) if data is not None else ""

        lines = []

        classification = data.get("classification", {})
        if isinstance(classification, dict):
            econ_type = classification.get("economic_type", "")
            sub_type = classification.get("sub_type", "")
            if econ_type:
                lines.append(f"**Тип экономики**: {econ_type}" + (f" / {sub_type}" if sub_type else ""))

        inventory = data.get("inventory", {})
        if isinstance(inventory, dict):
            resources = inventory.get("resources", [])
            if isinstance(resources, list):
                lines.append(f"\n**Ресурсы**: {len(resources)} типов")

        return "\n".join(lines) if lines else ""

    def _format_tech_tree(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование дерева технологий (разблокировок)."""
        if isinstance(data, dict):
            unlocks = data.get("unlockTree", data.get("unlocks", []))
            if isinstance(unlocks, list):
                return self._format_unlock_list(unlocks)
        elif isinstance(data, list):
            # Список разблокировок напрямую
            return self._format_unlock_list(data)
        return str(data) if data is not None else ""

    def _format_unlock_list(self, unlocks: list) -> str:
        """Форматирование списка разблокировок."""
        lines = ["### Дерево разблокировок"]
        for unlock in unlocks[:20]:
            if isinstance(unlock, dict):
                level = unlock.get("level", "?")
                name = unlock.get("unlock_name", unlock.get("name", ""))
                utype = unlock.get("unlock_type", "")
                lines.append(f"- Ур. {level}: **{name}** ({utype})")
            elif isinstance(unlock, str):
                lines.append(f"- {unlock}")
        return "\n".join(lines) if len(lines) > 1 else ""

    def _format_difficulty_curve(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование кривой сложности."""
        return self._format_difficulty(data, section_name, mapping)

    def _format_content_overview(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование обзора контента."""
        if not isinstance(data, dict):
            return str(data) if data is not None else ""

        lines = []
        content_plan = data.get("contentPlan", data)
        if isinstance(content_plan, dict):
            total = content_plan.get("total_content_requirements", {})
            if isinstance(total, dict):
                for k, v in total.items():
                    lines.append(f"- **{k}**: {v}")

        return "\n".join(lines) if lines else ""

    def _format_system_map(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование карты систем."""
        if not isinstance(data, dict):
            return str(data) if data is not None else ""

        lines = ["### Карта систем"]
        for key, value in data.items():
            if isinstance(value, (str, int, float)):
                lines.append(f"- **{key}**: {value}")
            elif isinstance(value, list):
                lines.append(f"- **{key}**: {len(value)} элементов")
            elif isinstance(value, dict):
                lines.append(f"- **{key}**: {len(value)} подсистем")

        return "\n".join(lines) if lines else ""

    def _format_feedback_patterns(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование паттернов обратной связи."""
        if not isinstance(data, dict):
            return str(data) if data is not None else ""

        lines = ["### Паттерны обратной связи"]

        loops = data.get("loops", data.get("inner_loops", []))
        if isinstance(loops, list):
            for loop in loops[:10]:
                if isinstance(loop, dict):
                    name = loop.get("name", "Петля")
                    ft = loop.get("feedback_type", loop.get("loop_type", ""))
                    lines.append(f"- **{name}** ({ft})")

        return "\n".join(lines) if lines else ""

    def _format_dissension(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование валидации лудонарративного диссонанса."""
        if not isinstance(data, dict):
            return str(data) if data is not None else ""

        lines = ["### Валидация лудонарративного диссонанса"]
        aesthetics = data.get("aesthetics", data.get("aesthetic_profile", {}))
        if isinstance(aesthetics, dict):
            primary = aesthetics.get("primary", "")
            if primary:
                lines.append(f"- **Целевая эстетика**: {primary}")

        return "\n".join(lines) if lines else ""

    def _format_pitch_deck(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование питч-дека."""
        if not isinstance(data, dict):
            return str(data) if data is not None else ""

        lines = ["### Pitch Deck"]
        for key in ["title", "genre", "description", "usp", "targetAudience"]:
            val = data.get(key)
            if val:
                if isinstance(val, (str, int, float)):
                    lines.append(f"- **{key}**: {val}")
                elif isinstance(val, list):
                    lines.append(f"- **{key}**: {', '.join(str(v) for v in val[:5])}")

        return "\n".join(lines) if lines else ""

    def _format_concept_overview(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование обзора концепции."""
        if not isinstance(data, dict):
            return str(data) if data is not None else ""

        lines = []
        for key in ["title", "genre", "description", "usp"]:
            val = data.get(key)
            if val and isinstance(val, str):
                label_map = {
                    "title": "Название",
                    "genre": "Жанр",
                    "description": "Описание",
                    "usp": "Уникальность",
                }
                lines.append(f"**{label_map.get(key, key)}**: {val}")

        return "\n\n".join(lines) if lines else ""

    def _format_mda_analysis(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Форматирование MDA-анализа."""
        if not isinstance(data, dict):
            return str(data) if data is not None else ""

        lines = ["### MDA-анализ"]

        aesthetics = data.get("aesthetics", data.get("aesthetic_profile", {}))
        if isinstance(aesthetics, dict):
            primary = aesthetics.get("primary", "")
            secondary = aesthetics.get("secondary", "")
            if primary:
                lines.append(f"- **Эстетика**: {primary}" + (f" / {secondary}" if secondary else ""))

        dynamics = data.get("dynamics_target", data.get("dynamics", {}))
        if isinstance(dynamics, dict):
            core = dynamics.get("core_dynamics", [])
            if isinstance(core, list) and core:
                lines.append(f"- **Динамики**: {', '.join(str(d) for d in core[:5])}")

        return "\n".join(lines) if lines else ""

    def _format_generic(
        self, data: Any, section_name: str, mapping: SectionMapping
    ) -> str:
        """Универсальное форматирование."""
        if isinstance(data, str):
            return data
        elif isinstance(data, dict):
            lines = []
            for k, v in data.items():
                if isinstance(v, (str, int, float, bool)):
                    lines.append(f"- **{k}**: {v}")
                elif isinstance(v, list):
                    lines.append(f"- **{k}**: {', '.join(str(i) for i in v[:10])}")
                elif isinstance(v, dict):
                    lines.append(f"- **{k}**: {len(v)} элементов")
            return "\n".join(lines) if lines else str(data)
        elif isinstance(data, list):
            lines = []
            for item in data[:15]:
                if isinstance(item, dict):
                    name = item.get("name", item.get("title", ""))
                    if name:
                        lines.append(f"- {name}")
                elif isinstance(item, str):
                    lines.append(f"- {item}")
            return "\n".join(lines) if lines else str(data)
        return str(data) if data is not None else ""

    # ========================================================
    # Вспомогательные методы для диаграмм/таблиц/формул
    # ========================================================

    def _generate_diagram_markdown(
        self, section_name: str, data: Any
    ) -> Optional[str]:
        """Генерация Markdown-представления диаграммы."""
        if not isinstance(data, dict):
            return None

        if section_name == "core_loop":
            steps = data.get("steps", [])
            if isinstance(steps, list) and steps:
                lines = ["```mermaid", "graph TD"]
                for i, step in enumerate(steps):
                    if isinstance(step, dict):
                        name = step.get("name", step.get("action", f"S{i+1}"))
                        safe_name = name.replace(" ", "_")[:20]
                        lines.append(f"    S{i+1}[{safe_name}]")
                        if i > 0:
                            lines.append(f"    S{i} --> S{i+1}")
                    else:
                        lines.append(f"    S{i+1}[{step}]")
                        if i > 0:
                            lines.append(f"    S{i} --> S{i+1}")
                # Замыкаем цикл
                if len(steps) > 1:
                    lines.append(f"    S{len(steps)} --> S1")
                lines.append("```")
                return "\n".join(lines)

        elif section_name in ("system_map", "feedback_patterns"):
            loops = data.get("loops", data.get("inner_loops", []))
            if isinstance(loops, list) and loops:
                lines = ["```mermaid", "graph LR"]
                for i, loop in enumerate(loops[:6]):
                    if isinstance(loop, dict):
                        name = loop.get("name", f"L{i+1}")
                        ft = loop.get("feedback_type", loop.get("loop_type", ""))
                        safe_name = name.replace(" ", "_")[:15]
                        color = "green" if ft in ("positive", "reinforcing") else "red"
                        lines.append(f"    L{i+1}({safe_name})")
                lines.append("```")
                return "\n".join(lines)

        elif section_name in ("progression", "progression_curves"):
            return "```mermaid\ngraph TD\n    XP[Опыт] --> Level[Уровень]\n    Level --> Power[Мощность]\n    Level --> Cost[Стоимость]\n```"

        elif section_name in ("economy", "economy_model"):
            return "```mermaid\ngraph LR\n    Source[Источники] --> Pool[Пул ресурсов]\n    Pool --> Drain[Стоки]\n    Pool --> Converter[Конвертер]\n    Converter --> Pool\n```"

        elif section_name in ("difficulty", "difficulty_curve"):
            return "```mermaid\ngraph TD\n    Base[Базовая сложность] --> Perceived[Воспринимаемая]\n    Tier[Tier Spike] --> Perceived\n```"

        return None

    def _generate_tables(
        self, section_name: str, data: Any
    ) -> Optional[list[dict]]:
        """Генерация таблиц для секции."""
        if not isinstance(data, dict):
            return None

        tables: list[dict] = []

        if section_name in ("mechanics",):
            mechanic_set = data.get("mechanicSet", data.get("mechanic_set", data))
            if isinstance(mechanic_set, dict):
                mechs = mechanic_set.get("mechanics", mechanic_set.get("base", []))
            elif isinstance(mechanic_set, list):
                mechs = mechanic_set
            else:
                mechs = []

            if isinstance(mechs, list) and mechs:
                rows = []
                for m in mechs[:15]:
                    if isinstance(m, dict):
                        rows.append({
                            "Название": m.get("name", m.get("mechanic_name", "")),
                            "Группа": m.get("group", m.get("group_name", "")),
                            "Описание": m.get("description", "")[:60],
                        })
                if rows:
                    tables.append({
                        "title": "Механики игры",
                        "headers": ["Название", "Группа", "Описание"],
                        "rows": rows,
                    })

        elif section_name in ("resources",):
            inventory = data.get("inventory", data)
            if isinstance(inventory, dict):
                resources = inventory.get("resources", [])
            elif isinstance(data, list):
                resources = data
            else:
                resources = []

            if isinstance(resources, list) and resources:
                rows = []
                for r in resources[:15]:
                    if isinstance(r, dict):
                        rows.append({
                            "Ресурс": r.get("name", ""),
                            "Класс": r.get("resource_class", ""),
                            "Потребляемый": "Да" if r.get("is_consumable") else "Нет",
                        })
                if rows:
                    tables.append({
                        "title": "Ресурсы игры",
                        "headers": ["Ресурс", "Класс", "Потребляемый"],
                        "rows": rows,
                    })

        elif section_name in ("balance", "balance_tables"):
            # Cost-power таблица
            objects = data.get("objects", data.get("balanceMap", {}).get("objects", []))
            if isinstance(objects, list) and objects:
                rows = []
                for obj in objects[:10]:
                    if isinstance(obj, dict):
                        rows.append({
                            "Объект": obj.get("name", ""),
                            "Cost": obj.get("cost", 0),
                            "Power": obj.get("power", 0),
                        })
                if rows:
                    tables.append({
                        "title": "Cost-Power баланс",
                        "headers": ["Объект", "Cost", "Power"],
                        "rows": rows,
                    })

        elif section_name in ("progression", "progression_curves"):
            curves = data.get("curves", {})
            if isinstance(curves, dict):
                xp = curves.get("xp_to_level", {})
                if isinstance(xp, dict):
                    tables.append({
                        "title": "Кривая XP → Level",
                        "headers": ["Параметр", "Значение"],
                        "rows": [
                            {"Параметр": k, "Значение": str(v)}
                            for k, v in xp.get("parameters", {}).items()
                            if isinstance(v, (int, float, str))
                        ],
                    })

        return tables if tables else None

    def _extract_formulas(
        self, section_name: str, data: Any
    ) -> Optional[list[str]]:
        """Извлечение формул из данных."""
        if not isinstance(data, dict):
            return None

        formulas: list[str] = []

        if section_name in ("progression", "progression_curves", "difficulty_curve", "difficulty"):
            curves = data.get("curves", data)
            if isinstance(curves, dict):
                for curve_key in ["xp_to_level", "level_to_power", "level_to_cost", "difficulty"]:
                    curve = curves.get(curve_key, {})
                    if isinstance(curve, dict) and curve.get("formula"):
                        formulas.append(f"{curve_key}: {curve['formula']}")

        elif section_name in ("balance", "balance_tables"):
            formula = data.get("cost_power_formula", data.get("formula", ""))
            if formula and isinstance(formula, str):
                formulas.append(f"cost/power: {formula}")

        elif section_name in ("economy", "economy_model"):
            classification = data.get("classification", {})
            if isinstance(classification, dict):
                econ_type = classification.get("economic_type", "")
                formulas.append(f"Тип экономики: {econ_type}")

        return formulas if formulas else None

    # ========================================================
    # Полный пайплайн: Этапы 1–3
    # ========================================================

    async def generate_stages_1_3(
        self,
        input_data: GDDGenerationInput,
    ) -> GDDProfile:
        """
        Полный пайплайн генерации GDD — Этапы 1–3 алгоритма 3.7.

        Выполняет последовательно:
        1. Определение формата GDD
        2. Маппинг Project State → секции
        3. Автозаполнение секций

        Returns:
            GDDProfile с результатами всех трёх этапов
        """
        pipeline_start = time.time()

        # === Этап 1: Формат ===
        format_spec = await self.determine_gdd_format(input_data)

        # === Этап 2: Маппинг ===
        data_mapping = await self.map_project_to_sections(
            format_spec, input_data
        )

        # === Этап 3: Автозаполнение ===
        auto_filled = await self.generate_auto_sections(
            data_mapping, input_data
        )

        latency_ms = int((time.time() - pipeline_start) * 1000)

        # Общий coverage_score
        coverage_score = data_mapping.coverage_score

        profile = GDDProfile(
            format_spec=format_spec,
            data_mapping=data_mapping,
            auto_filled_sections=auto_filled,
            stages_completed=[1, 2, 3],
            coverage_score=coverage_score,
            latency_ms=latency_ms,
        )

        logger.info(
            f"[Pipeline 1-3] Completed in {latency_ms}ms. "
            f"Format: {format_spec.format}, "
            f"Sections: {len(format_spec.sections)}, "
            f"Auto-filled: {auto_filled.count}, "
            f"Coverage: {coverage_score:.2f}"
        )

        return profile

    # ========================================================
    # Этап 4: AI-генерация и обогащение секций (3.7.6)
    # ========================================================

    # Маппинг секций → AI-промпт для генерации
    SECTION_PROMPT_MAP: dict[str, str] = {
        "characters": "GENERATE_CHARACTERS_SECTION",
        "character_profiles": "GENERATE_CHARACTERS_SECTION",
        "story": "GENERATE_STORY_SECTION",
        "dialogues": "GENERATE_STORY_SECTION",
        "quests": "GENERATE_STORY_SECTION",
        "quest_matrix": "GENERATE_STORY_SECTION",
        "lore": "GENERATE_STORY_SECTION",
        "lore_db": "GENERATE_STORY_SECTION",
        "storyline_map": "GENERATE_STORY_SECTION",
        "cutscene_scripts": "GENERATE_STORY_SECTION",
        "narrative_arc": "GENERATE_STORY_SECTION",
        "visual_style": "GENERATE_VISUAL_STYLE",
        "visual_hook": "GENERATE_VISUAL_STYLE",
        "controls": "GENERATE_CONTROLS_SECTION",
        "camera_perspective": "GENERATE_CONTROLS_SECTION",
        "world_structure": "GENERATE_WORLD_SECTION",
        "level_design": "GENERATE_WORLD_SECTION",
        "navigation": "GENERATE_WORLD_SECTION",
        "combat_spaces": "GENERATE_WORLD_SECTION",
        "sound_music": "GENERATE_AUDIO_SECTION",
    }

    async def generate_ai_sections(
        self,
        data_mapping: GDDDataMapping,
        auto_filled: AutoFilledSections,
        input_data: GDDGenerationInput,
    ) -> AIEnrichedSections:
        """
        Этап 4: AI-генерация и обогащение секций GDD.

        Алгоритм 3.7.6:
        1. Для секций с ai_enrich=True, которые были автозаполнены:
           - Вызвать ENRICH_SECTION для обогащения контекста
        2. Для секций с ai_generate=True:
           - Вызвать специализированный промпт генерации
        3. Обработать ошибки gracefully (добавить в failed_sections)
        4. Рассчитать total_coverage после AI-обогащения

        Returns:
            AIEnrichedSections с обогащёнными и сгенерированными секциями
        """
        start = time.time()

        enriched_sections: dict[str, SectionContent] = {}
        generated_sections: dict[str, SectionContent] = {}
        failed_sections: list[str] = []

        # === Шаг 4.1: Обогащение автозаполненных секций ===
        for section_name, section_content in auto_filled.sections.items():
            mapping = data_mapping.active_mappings.get(section_name)
            if not mapping or not mapping.ai_enrich:
                continue

            try:
                enriched_content = await self._enrich_section(
                    section_name, section_content, input_data
                )
                if enriched_content:
                    enriched_sections[section_name] = enriched_content
                    logger.debug(f"[Stage 4.1] Enriched: {section_name}")
                else:
                    # Обогащение не удалось, оставляем оригинал
                    enriched_sections[section_name] = section_content
            except Exception as e:
                logger.warning(
                    f"[Stage 4.1] Failed to enrich '{section_name}': {e}"
                )
                failed_sections.append(section_name)
                # Оставляем оригинальный контент
                enriched_sections[section_name] = section_content

        # === Шаг 4.2: Генерация секций с нуля ===
        ai_generatable = data_mapping.ai_generatable_sections
        for section_name in ai_generatable:
            mapping = data_mapping.active_mappings.get(section_name)
            if not mapping or not mapping.ai_generate:
                continue

            # Пропускаем секции, которые уже были автозаполнены и обогащены
            if section_name in auto_filled.sections:
                continue

            try:
                generated_content = await self._generate_section(
                    section_name, mapping, data_mapping, input_data
                )
                if generated_content:
                    generated_sections[section_name] = generated_content
                    logger.debug(f"[Stage 4.2] Generated: {section_name}")
                else:
                    failed_sections.append(section_name)
            except Exception as e:
                logger.warning(
                    f"[Stage 4.2] Failed to generate '{section_name}': {e}"
                )
                failed_sections.append(section_name)

        # === Шаг 4.3: Рассчитываем total_coverage ===
        total_sections = len(data_mapping.format_spec.sections)
        filled_count = len(auto_filled.sections) + len(generated_sections)
        total_coverage = round(filled_count / total_sections, 3) if total_sections > 0 else 0.0

        result = AIEnrichedSections(
            enriched_sections=enriched_sections,
            generated_sections=generated_sections,
            enriched_count=len(enriched_sections),
            generated_count=len(generated_sections),
            failed_sections=failed_sections,
            total_coverage=total_coverage,
        )

        logger.info(
            f"[Stage 4] AI processing: enriched={len(enriched_sections)}, "
            f"generated={len(generated_sections)}, "
            f"failed={len(failed_sections)}, "
            f"total_coverage={total_coverage:.2f} "
            f"({time.time() - start:.2f}s)"
        )

        return result

    async def _enrich_section(
        self,
        section_name: str,
        section_content: SectionContent,
        input_data: GDDGenerationInput,
    ) -> Optional[SectionContent]:
        """
        Обогащение автозаполненной секции через ENRICH_SECTION промпт.
        """
        genre = input_data.concept.get("genre", "") if input_data.concept else ""
        aesthetics = (
            input_data.mda_profile.get("aesthetics", {})
            if input_data.mda_profile else {}
        )

        prompt_inputs = {
            "section_name": section_name,
            "current_content": section_content.content[:2000],
            "genre": genre,
            "aesthetics": json.dumps(aesthetics, ensure_ascii=False) if aesthetics else "",
        }

        result = await self.executor.execute(
            prompt_id="ENRICH_SECTION",
            inputs=prompt_inputs,
        )

        enriched_text = self._extract_text_from_result(result.data)

        if not enriched_text:
            return None

        return SectionContent(
            content=enriched_text,
            source="ai_enrich",
            auto_filled=True,
            diagram=section_content.diagram,
            tables=section_content.tables,
            formulas=section_content.formulas,
            requires_review=True,
        )

    async def _generate_section(
        self,
        section_name: str,
        mapping: SectionMapping,
        data_mapping: GDDDataMapping,
        input_data: GDDGenerationInput,
    ) -> Optional[SectionContent]:
        """
        Генерация секции с нуля через специализированный AI-промпт.
        """
        # Определяем промпт по секции
        prompt_id = self.SECTION_PROMPT_MAP.get(section_name, "ENRICH_SECTION")

        # Собираем входные данные для промпта
        prompt_inputs = self._build_generation_inputs(
            section_name, prompt_id, input_data
        )

        result = await self.executor.execute(
            prompt_id=prompt_id,
            inputs=prompt_inputs,
        )

        generated_text = self._extract_text_from_result(result.data)

        if not generated_text:
            return None

        # Определяем дополнительные элементы по маппингу
        diagram = None
        tables = None
        formulas = None

        if mapping.diagram:
            diagram = "```mermaid\ngraph TD\n    Generated[AI Generated Section]\n```"

        if mapping.tables:
            # Пробуем извлечь таблицы из результата
            tables = self._try_extract_tables(result.data)

        return SectionContent(
            content=generated_text,
            source="ai_generate",
            auto_filled=False,
            diagram=diagram,
            tables=tables,
            formulas=formulas,
            requires_review=True,
        )

    def _build_generation_inputs(
        self,
        section_name: str,
        prompt_id: str,
        input_data: GDDGenerationInput,
    ) -> dict[str, Any]:
        """Сборка входных данных для промпта генерации секции."""
        genre = input_data.concept.get("genre", "") if input_data.concept else ""
        platforms = (
            input_data.concept.get("platforms", [])
            if input_data.concept else []
        )
        core_loop = input_data.core_loop or {}
        mechanics = (
            input_data.mda_profile.get("mechanicSet", {}).get("mechanics", [])
            if input_data.mda_profile else []
        )
        aesthetics = (
            input_data.mda_profile.get("aesthetics", {})
            if input_data.mda_profile else {}
        )

        # Базовые входы для большинства промптов
        base_inputs: dict[str, Any] = {
            "genre": genre,
            "platforms": platforms if isinstance(platforms, list) else [platforms],
            "core_loop": core_loop,
            "mechanics": mechanics,
            "aesthetics": aesthetics,
            "section_name": section_name,
        }

        # Специфичные входы для конкретных промптов
        if prompt_id == "GENERATE_CHARACTERS_SECTION":
            base_inputs["setting"] = (
                input_data.concept.get("description", "")
                if input_data.concept else ""
            )
            base_inputs["mood"] = (
                input_data.mda_profile.get("aesthetics", {}).get("primary", "")
                if input_data.mda_profile else ""
            )

        elif prompt_id == "GENERATE_VISUAL_STYLE":
            base_inputs["mood"] = (
                input_data.mda_profile.get("aesthetics", {}).get("primary", "")
                if input_data.mda_profile else ""
            )
            base_inputs["setting"] = (
                input_data.concept.get("description", "")
                if input_data.concept else ""
            )

        elif prompt_id == "GENERATE_STORY_SECTION":
            base_inputs["setting"] = (
                input_data.concept.get("description", "")
                if input_data.concept else ""
            )
            base_inputs["usp"] = (
                input_data.concept.get("usp", "")
                if input_data.concept else ""
            )

        elif prompt_id == "GENERATE_CONTROLS_SECTION":
            base_inputs["camera_perspective"] = ""  # Нет данных по умолчанию

        elif prompt_id == "GENERATE_WORLD_SECTION":
            base_inputs["setting"] = (
                input_data.concept.get("description", "")
                if input_data.concept else ""
            )
            base_inputs["progression"] = input_data.progression_profile or {}

        elif prompt_id == "GENERATE_AUDIO_SECTION":
            base_inputs["mood"] = (
                input_data.mda_profile.get("aesthetics", {}).get("primary", "")
                if input_data.mda_profile else ""
            )
            base_inputs["visual_style"] = ""

        return base_inputs

    def _extract_text_from_result(self, data: Any) -> str:
        """Извлечение текста из результата AI-вызова."""
        if isinstance(data, str):
            return data
        elif isinstance(data, dict):
            # Пробуем типичные ключи с текстом
            text_keys = [
                "content", "text", "description", "markdown", "section_content",
                "enriched_content", "generated_content",
            ]
            for key in text_keys:
                if key in data and isinstance(data[key], str):
                    return data[key]

            # Если есть вложенные объекты — форматируем как Markdown
            lines = []
            for k, v in data.items():
                if isinstance(v, str):
                    if len(v) > 50:
                        lines.append(f"## {k}\n{v}")
                    else:
                        lines.append(f"**{k}**: {v}")
                elif isinstance(v, list):
                    lines.append(f"## {k}")
                    for item in v[:15]:
                        if isinstance(item, str):
                            lines.append(f"- {item}")
                        elif isinstance(item, dict):
                            name = item.get("name", item.get("title", ""))
                            if name:
                                lines.append(f"- **{name}**")
                            else:
                                lines.append(f"- {json.dumps(item, ensure_ascii=False)[:100]}")
                elif isinstance(v, dict):
                    lines.append(f"## {k}")
                    for dk, dv in v.items():
                        if isinstance(dv, (str, int, float)):
                            lines.append(f"- **{dk}**: {dv}")

            return "\n\n".join(lines) if lines else json.dumps(data, ensure_ascii=False, indent=2)
        elif isinstance(data, list):
            lines = []
            for item in data[:20]:
                if isinstance(item, str):
                    lines.append(f"- {item}")
                elif isinstance(item, dict):
                    name = item.get("name", item.get("title", ""))
                    desc = item.get("description", "")
                    if name:
                        line = f"- **{name}**"
                        if desc:
                            line += f": {desc[:100]}"
                        lines.append(line)
            return "\n".join(lines) if lines else json.dumps(data, ensure_ascii=False, indent=2)

        return str(data) if data is not None else ""

    def _try_extract_tables(self, data: Any) -> Optional[list[dict]]:
        """Попытка извлечь табличные данные из AI-результата."""
        if isinstance(data, dict):
            table_keys = ["tables", "data", "items", "entries", "list"]
            for key in table_keys:
                val = data.get(key)
                if isinstance(val, list) and val:
                    rows = []
                    for item in val[:15]:
                        if isinstance(item, dict):
                            rows.append(item)
                    if rows:
                        headers = list(rows[0].keys())[:5]
                        return [{
                            "title": key.capitalize(),
                            "headers": headers,
                            "rows": rows,
                        }]
        return None

    # ========================================================
    # Этап 5: Ручные секции с подсказками (3.7.7)
    # ========================================================

    # Классификация секций по приоритету
    CRITICAL_SECTIONS: set[str] = {
        "core_loop", "mechanics", "progression", "balance",
        "resources", "economy",
    }

    IMPORTANT_SECTIONS: set[str] = {
        "characters", "story", "world_structure", "level_design",
        "hud_ui", "visual_style",
    }

    OPTIONAL_SECTIONS: set[str] = {
        "controls", "dialogues", "quests", "lore", "sound_music",
        "multiplayer_modes", "social_features", "meta_game",
        "tech_requirements", "milestones_budget", "menus",
        "camera_perspective", "navigation", "combat_spaces",
        "game_modes", "platform_ports", "monetization",
    }

    # Шаблоны-скелеты для ручных секций
    SECTION_TEMPLATES: dict[str, str] = {
        "controls": (
            "## Управление\n\n"
            "### Основная схема\n"
            "[Опишите основную схему управления]\n\n"
            "### Платформо-зависимые маппинги\n"
            "- **PC (Keyboard+Mouse)**: [маппинг]\n"
            "- **Gamepad**: [маппинг]\n"
            "- **Touch**: [маппинг]\n\n"
            "### Ключевые действия\n"
            "| Действие | PC | Gamepad |\n|---|---|---|\n"
            "| [действие] | [клавиша] | [кнопка] |"
        ),
        "characters": (
            "## Персонажи\n\n"
            "### Главный герой\n"
            "- **Имя**: [имя]\n"
            "- **Роль**: [роль в истории]\n"
            "- **Мотивация**: [почему действует]\n\n"
            "### Ключевые NPC\n"
            "| Имя | Роль | Отношение к герою |\n|---|---|---|\n"
            "| [имя] | [роль] | [отношение] |"
        ),
        "story": (
            "## Сюжет\n\n"
            "### Логлайн\n"
            "[1-2 предложения сути истории]\n\n"
            "### Акт 1 — Завязка\n"
            "[Начальная ситуация, зов к приключению]\n\n"
            "### Акт 2 — Конфликт\n"
            "[Развитие конфликта, препятствия]\n\n"
            "### Акт 3 — Разрешение\n"
            "[Кульминация и развязка]"
        ),
        "dialogues": (
            "## Диалоги\n\n"
            "### Принципы\n"
            "- [Принцип 1: например, диалог раскрывает характер]\n"
            "- [Принцип 2: диалог двигает сюжет]\n\n"
            "### Пример диалога\n"
            "```\n[Персонаж А]: [реплика]\n[Персонаж Б]: [реплика]\n```"
        ),
        "quests": (
            "## Квесты\n\n"
            "### Основной квест\n"
            "| ID | Название | Описание | Награда |\n|---|---|---|---|\n"
            "| MQ01 | [название] | [описание] | [награда] |\n\n"
            "### Побочные квесты\n"
            "| ID | Название | Связь с NPC | Награда |\n|---|---|---|---|\n"
            "| SQ01 | [название] | [NPC] | [награда] |"
        ),
        "lore": (
            "## Лор\n\n"
            "### Мир\n"
            "[Описание мира и его правил]\n\n"
            "### История\n"
            "[Ключевые события прошлого]\n\n"
            "### Фракции\n"
            "| Название | Идеология | Отношение к герою |\n|---|---|---|\n"
            "| [название] | [идеология] | [отношение] |"
        ),
        "world_structure": (
            "## Структура мира\n\n"
            "### Общая карта\n"
            "[Описание общей структуры мира]\n\n"
            "### Локации\n"
            "| Локация | Размер | Ключевые объекты | Уровень опасности |\n|---|---|---|---|\n"
            "| [название] | [размер] | [объекты] | [уровень] |"
        ),
        "level_design": (
            "## Дизайн уровней\n\n"
            "### Принципы\n"
            "- [Принцип 1]\n"
            "- [Принцип 2]\n\n"
            "### Шаблон уровня\n"
            "[Описание типичной структуры уровня]"
        ),
        "navigation": (
            "## Навигация\n\n"
            "### Способы перемещения\n"
            "- [Способ 1]\n"
            "- [Способ 2]\n\n"
            "### Карта мира\n"
            "[Описание навигации по карте]"
        ),
        "combat_spaces": (
            "## Боевые пространства\n\n"
            "### Принципы дизайна\n"
            "- [Принцип 1]\n"
            "- [Принцип 2]\n\n"
            "### Типы арен\n"
            "| Тип | Размер | Особенности |\n|---|---|---|\n"
            "| [тип] | [размер] | [особенности] |"
        ),
        "visual_style": (
            "## Визуальный стиль\n\n"
            "### Арт-направление\n"
            "[Описание визуального стиля]\n\n"
            "### Цветовая палитра\n"
            "[Основные цвета и их назначение]\n\n"
            "### Референсы\n"
            "- [Референс 1]\n"
            "- [Референс 2]"
        ),
        "sound_music": (
            "## Звук и музыка\n\n"
            "### Музыкальный стиль\n"
            "[Описание]\n\n"
            "### Приоритетные SFX\n"
            "| Категория | Приоритет | Описание |\n|---|---|---|\n"
            "| [категория] | [приоритет] | [описание] |"
        ),
        "multiplayer_modes": (
            "## Мультиплеер\n\n"
            "### Режимы\n"
            "| Режим | Игроков | Описание |\n|---|---|---|\n"
            "| [режим] | [число] | [описание] |"
        ),
        "social_features": (
            "## Социальные функции\n\n"
            "### Список функций\n"
            "- [Функция 1]\n"
            "- [Функция 2]"
        ),
        "meta_game": (
            "## Мета-игра\n\n"
            "### Сезонный контент\n"
            "[Описание подхода к сезонам]\n\n"
            "### Эндгейм\n"
            "[Описание эндгейм-контента]"
        ),
        "tech_requirements": (
            "## Технические требования\n\n"
            "### Минимальные\n"
            "- **CPU**: [требование]\n"
            "- **GPU**: [требование]\n"
            "- **RAM**: [требование]\n\n"
            "### Рекомендуемые\n"
            "- **CPU**: [требование]\n"
            "- **GPU**: [требование]\n"
            "- **RAM**: [требование]"
        ),
        "milestones_budget": (
            "## Майлстоуны и бюджет\n\n"
            "| Майлстоун | Срок | Бюджет | Результат |\n|---|---|---|---|\n"
            "| [M1] | [срок] | [бюджет] | [результат] |"
        ),
        "monetization": (
            "## Монетизация\n\n"
            "### Модель\n"
            "[Описание модели монетизации]\n\n"
            "### Ценовые точки\n"
            "| Товар | Цена | Ценность для игрока |\n|---|---|---|\n"
            "| [товар] | [цена] | [ценность] |"
        ),
        "menus": (
            "## Меню и UI\n\n"
            "### Главное меню\n"
            "[Описание структуры главного меню]\n\n"
            "### In-game HUD\n"
            "[Описание HUD]"
        ),
        "hud_ui": (
            "## HUD и UI\n\n"
            "### Элементы HUD\n"
            "| Элемент | Позиция | Когда виден |\n|---|---|---|\n"
            "| [элемент] | [позиция] | [условие] |"
        ),
        "camera_perspective": (
            "## Камера\n\n"
            "### Перспектива\n"
            "[Описание типа камеры]\n\n"
            "### Поведение\n"
            "[Как камера реагирует на действия игрока]"
        ),
        "game_modes": (
            "## Игровые режимы\n\n"
            "| Режим | Описание | Длительность сессии |\n|---|---|---|\n"
            "| [режим] | [описание] | [длительность] |"
        ),
    }

    async def generate_manual_skeletons(
        self,
        data_mapping: GDDDataMapping,
        input_data: GDDGenerationInput,
    ) -> ManualSectionsResult:
        """
        Этап 5: Генерация скелетов ручных секций с AI-подсказками.

        Алгоритм 3.7.7:
        1. Для каждой секции из manual_sections:
           - Генерировать шаблон-скелет с плейсхолдерами
           - Вызвать AI_GENERATE_SECTION_HINTS для подсказок
           - Классифицировать по приоритету: critical/important/optional
        2. Обработать ошибки gracefully

        Returns:
            ManualSectionsResult со скелетами секций
        """
        start = time.time()

        skeletons: dict[str, ManualSectionSkeleton] = {}
        critical_sections: list[str] = []
        important_sections: list[str] = []
        optional_sections: list[str] = []
        failed_sections: list[str] = []

        manual_sections = data_mapping.manual_sections

        for section_name in manual_sections:
            mapping = data_mapping.active_mappings.get(section_name)

            # Определяем приоритет
            priority = self._classify_section_priority(section_name)

            if priority == "critical":
                critical_sections.append(section_name)
            elif priority == "important":
                important_sections.append(section_name)
            else:
                optional_sections.append(section_name)

            # Генерируем шаблон-скелет
            template = self._generate_section_template(section_name)

            # Оцениваем трудоёмкость
            estimated_effort = self._estimate_effort(section_name, mapping)

            # Получаем AI-подсказки
            hints: list[str] = []
            try:
                hints = await self._generate_section_hints(
                    section_name, input_data
                )
            except Exception as e:
                logger.warning(
                    f"[Stage 5] Failed to get hints for '{section_name}': {e}"
                )
                failed_sections.append(section_name)

            skeleton = ManualSectionSkeleton(
                section_name=section_name,
                priority=priority,
                template=template,
                hints=hints,
                estimated_effort=estimated_effort,
            )
            skeletons[section_name] = skeleton

        total_manual_count = len(skeletons)

        result = ManualSectionsResult(
            skeletons=skeletons,
            critical_sections=critical_sections,
            important_sections=important_sections,
            optional_sections=optional_sections,
            total_manual_count=total_manual_count,
            failed_sections=failed_sections,
        )

        logger.info(
            f"[Stage 5] Manual skeletons: {total_manual_count} sections, "
            f"critical={len(critical_sections)}, "
            f"important={len(important_sections)}, "
            f"optional={len(optional_sections)}, "
            f"failed={len(failed_sections)} "
            f"({time.time() - start:.2f}s)"
        )

        return result

    def _classify_section_priority(self, section_name: str) -> SectionPriority:
        """Классификация секции по приоритету."""
        if section_name in self.CRITICAL_SECTIONS:
            return "critical"
        elif section_name in self.IMPORTANT_SECTIONS:
            return "important"
        else:
            return "optional"

    def _generate_section_template(self, section_name: str) -> str:
        """Генерация шаблона-скелета для ручной секции."""
        if section_name in self.SECTION_TEMPLATES:
            return self.SECTION_TEMPLATES[section_name]

        # Универсальный шаблон для неизвестных секций
        return (
            f"## {section_name.replace('_', ' ').title()}\n\n"
            "### Описание\n"
            "[Опишите раздел]\n\n"
            "### Ключевые элементы\n"
            "- [Элемент 1]\n"
            "- [Элемент 2]\n"
            "- [Элемент 3]\n\n"
            "### Примечания\n"
            "[Дополнительные заметки]"
        )

    def _estimate_effort(
        self, section_name: str, mapping: Optional[SectionMapping]
    ) -> str:
        """Оценка трудоёмкости заполнения секции."""
        if section_name in self.CRITICAL_SECTIONS:
            return "high"
        elif section_name in self.IMPORTANT_SECTIONS:
            return "medium"
        elif mapping and (mapping.diagram or mapping.formulas):
            return "high"
        elif mapping and mapping.tables:
            return "medium"
        else:
            return "low"

    async def _generate_section_hints(
        self,
        section_name: str,
        input_data: GDDGenerationInput,
    ) -> list[str]:
        """Получение AI-подсказок для заполнения ручной секции."""
        genre = input_data.concept.get("genre", "") if input_data.concept else ""
        description = (
            input_data.concept.get("description", "")
            if input_data.concept else ""
        )

        prompt_inputs = {
            "section": section_name,
            "genre": genre,
            "project_summary": description[:500] if description else "",
        }

        result = await self.executor.execute(
            prompt_id="AI_GENERATE_SECTION_HINTS",
            inputs=prompt_inputs,
        )

        # Извлекаем подсказки из результата
        hints = self._extract_hints_from_result(result.data)
        return hints

    def _extract_hints_from_result(self, data: Any) -> list[str]:
        """Извлечение подсказок из результата AI."""
        hints: list[str] = []

        if isinstance(data, dict):
            # Ищем список подсказок
            hint_keys = ["hints", "suggestions", "tips", "recommendations", "items"]
            for key in hint_keys:
                val = data.get(key)
                if isinstance(val, list):
                    for item in val[:10]:
                        if isinstance(item, str):
                            hints.append(item)
                        elif isinstance(item, dict):
                            text = item.get("hint", item.get("text", item.get("suggestion", "")))
                            if text:
                                hints.append(str(text))
                    break

            if not hints:
                # Пробуем создать подсказки из полей
                for k, v in data.items():
                    if isinstance(v, str) and len(v) > 20:
                        hints.append(f"**{k}**: {v[:200]}")
                    elif isinstance(v, list) and len(v) > 0:
                        first = v[0]
                        if isinstance(first, str):
                            hints.append(f"Рассмотрите: {', '.join(str(i) for i in v[:5])}")
                        break

        elif isinstance(data, list):
            for item in data[:10]:
                if isinstance(item, str):
                    hints.append(item)
                elif isinstance(item, dict):
                    text = item.get("hint", item.get("text", item.get("suggestion", "")))
                    if text:
                        hints.append(str(text))

        elif isinstance(data, str):
            # Разбиваем строку на подсказки
            lines = [line.strip() for line in data.split("\n") if line.strip()]
            hints = lines[:10]

        return hints[:10] if hints else [
            "Опишите ключевые элементы раздела",
            "Укажите связи с другими секциями GDD",
            "Добавьте примеры и конкретные значения",
        ]

    # ========================================================
    # Полный пайплайн: Этапы 1–5
    # ========================================================

    async def generate_stages_1_5(
        self,
        input_data: GDDGenerationInput,
    ) -> GDDProfile:
        """
        Полный пайплайн генерации GDD — Этапы 1–5 алгоритма 3.7.

        Выполняет последовательно:
        1. Определение формата GDD
        2. Маппинг Project State → секции
        3. Автозаполнение секций
        4. AI-генерация и обогащение секций
        5. Генерация скелетов ручных секций с подсказками

        Returns:
            GDDProfile с результатами всех пяти этапов
        """
        pipeline_start = time.time()

        # === Этап 1: Формат ===
        format_spec = await self.determine_gdd_format(input_data)

        # === Этап 2: Маппинг ===
        data_mapping = await self.map_project_to_sections(
            format_spec, input_data
        )

        # === Этап 3: Автозаполнение ===
        auto_filled = await self.generate_auto_sections(
            data_mapping, input_data
        )

        # === Этап 4: AI-генерация и обогащение ===
        ai_enriched = await self.generate_ai_sections(
            data_mapping, auto_filled, input_data
        )

        # === Этап 5: Ручные секции ===
        manual_skeletons = await self.generate_manual_skeletons(
            data_mapping, input_data
        )

        latency_ms = int((time.time() - pipeline_start) * 1000)

        # Общий coverage_score: учитываем все заполненные секции
        total_sections = len(format_spec.sections)
        all_filled = (
            len(auto_filled.sections)
            + ai_enriched.generated_count
        )
        coverage_score = round(all_filled / total_sections, 3) if total_sections > 0 else 0.0

        profile = GDDProfile(
            format_spec=format_spec,
            data_mapping=data_mapping,
            auto_filled_sections=auto_filled,
            ai_enriched_sections=ai_enriched,
            manual_skeletons=manual_skeletons,
            stages_completed=[1, 2, 3, 4, 5],
            coverage_score=coverage_score,
            latency_ms=latency_ms,
        )

        logger.info(
            f"[Pipeline 1-5] Completed in {latency_ms}ms. "
            f"Format: {format_spec.format}, "
            f"Sections: {len(format_spec.sections)}, "
            f"Auto-filled: {auto_filled.count}, "
            f"AI-enriched: {ai_enriched.enriched_count}, "
            f"AI-generated: {ai_enriched.generated_count}, "
            f"Manual: {manual_skeletons.total_manual_count}, "
            f"Coverage: {coverage_score:.2f}"
        )

        return profile

    # ========================================================
    # Этап 6: Сшивка и валидация GDD (3.7.8)
    # ========================================================

    async def assemble_gdd(self, profile: GDDProfile) -> GDDAssembledDocument:
        """
        Этап 6: Сшивка и валидация GDD (алгоритм 3.7.8).

        Объединяет все секции из этапов 3-5 в единый документ,
        проверяет согласованность между секциями, обнаруживает
        лудонарративный диссонанс.
        """
        start = time.time()

        assembled: dict[str, GDDAssembledSection] = {}
        format_spec = profile.format_spec
        all_sections = format_spec.sections

        # Собираем контент из всех источников с приоритетом:
        # ai_enriched > auto_filled > manual_skeletons
        for section_name in all_sections:
            content_text = ""
            source = ""
            has_diagram = False
            has_tables = False
            has_formulas = False
            requires_review = False

            # Приоритет 1: AI-обогащённые секции (enriched)
            if profile.ai_enriched_sections:
                enriched = profile.ai_enriched_sections.enriched_sections.get(section_name)
                if enriched and enriched.content:
                    content_text = enriched.content
                    source = "ai_enrich"
                    has_diagram = enriched.diagram is not None
                    has_tables = enriched.tables is not None and len(enriched.tables) > 0
                    has_formulas = enriched.formulas is not None and len(enriched.formulas) > 0
                    requires_review = enriched.requires_review

            # Приоритет 2: AI-сгенерированные секции
            if not content_text and profile.ai_enriched_sections:
                generated = profile.ai_enriched_sections.generated_sections.get(section_name)
                if generated and generated.content:
                    content_text = generated.content
                    source = "ai_generate"
                    has_diagram = generated.diagram is not None
                    has_tables = generated.tables is not None and len(generated.tables) > 0
                    has_formulas = generated.formulas is not None and len(generated.formulas) > 0
                    requires_review = generated.requires_review

            # Приоритет 3: Автозаполненные секции
            if not content_text and profile.auto_filled_sections:
                auto_section = profile.auto_filled_sections.sections.get(section_name)
                if auto_section and auto_section.content:
                    content_text = auto_section.content
                    source = "auto_fill"
                    has_diagram = auto_section.diagram is not None
                    has_tables = auto_section.tables is not None and len(auto_section.tables) > 0
                    has_formulas = auto_section.formulas is not None and len(auto_section.formulas) > 0
                    requires_review = auto_section.requires_review

            # Приоритет 4: Скелеты ручных секций (шаблонный контент)
            if not content_text and profile.manual_skeletons:
                skeleton = profile.manual_skeletons.skeletons.get(section_name)
                if skeleton and skeleton.template:
                    content_text = skeleton.template
                    source = "manual"
                    requires_review = True

            assembled[section_name] = GDDAssembledSection(
                section_name=section_name,
                content=content_text,
                source=source,
                has_diagram=has_diagram,
                has_tables=has_tables,
                has_formulas=has_formulas,
                requires_review=requires_review,
            )

        # Валидация согласованности
        input_data = self._reconstruct_input_from_profile(profile)
        consistency_report = self.validate_consistency(assembled, input_data)

        # Порядок секций из format_spec
        section_order = list(format_spec.sections)

        # Считаем покрытие
        total_sections = len(all_sections)
        filled_sections = sum(
            1 for s in assembled.values() if s.content and s.source not in ("", "manual")
        )
        coverage_score = round(filled_sections / total_sections, 3) if total_sections > 0 else 0.0

        result = GDDAssembledDocument(
            sections=assembled,
            section_order=section_order,
            consistency_report=consistency_report,
            total_sections=total_sections,
            filled_sections=filled_sections,
            coverage_score=coverage_score,
        )

        logger.info(
            f"[Stage 6] Assembled: {filled_sections}/{total_sections} sections filled, "
            f"coverage={coverage_score:.2f}, "
            f"consistency: {consistency_report.error_count} errors, "
            f"{consistency_report.warning_count} warnings, "
            f"{consistency_report.info_count} info "
            f"({time.time() - start:.2f}s)"
        )

        return result

    def validate_consistency(
        self,
        assembled: dict[str, GDDAssembledSection],
        input_data: GDDGenerationInput,
    ) -> ConsistencyReport:
        """
        Валидация согласованности между секциями GDD.

        Проверяет пары секций на предмет:
        - Core Loop ↔ Mechanics: шаги core loop ссылаются на механики
        - Progression ↔ Balance: кривые прогрессии согласованы с балансом
        - Economy ↔ Monetization: ресурсы экономики поддерживают монетизацию
        - Narrative ↔ Mechanics: нарратив согласован с механиками
        - Лудонарративный диссонанс: тон нарратива не противоречит геймплею

        Детерминистическая проверка (без AI-вызовов).
        """
        issues: list[ConsistencyIssue] = []
        checked_pairs: list[str] = []

        def _has_content(name: str) -> bool:
            sec = assembled.get(name)
            return sec is not None and bool(sec.content) and sec.source != ""

        def _get_content(name: str) -> str:
            sec = assembled.get(name)
            return sec.content if sec else ""

        # ── Пара 1: Core Loop ↔ Mechanics ──
        pair_key = "core_loop::mechanics"
        checked_pairs.append(pair_key)

        core_loop_present = _has_content("core_loop")
        mechanics_present = _has_content("mechanics")

        if core_loop_present and mechanics_present:
            core_content = _get_content("core_loop").lower()
            mech_content = _get_content("mechanics").lower()

            # Проверяем: есть ли в core loop упоминания механик
            mech_names = re.findall(r'\*\*([^*]+)\*\*', _get_content("mechanics"))
            referenced_mechs = sum(
                1 for name in mech_names
                if name.lower() in core_content
            )

            if referenced_mechs == 0 and len(mech_names) > 0:
                issues.append(ConsistencyIssue(
                    severity="warning",
                    section_a="core_loop",
                    section_b="mechanics",
                    issue_type="coreloop_mechanics_mismatch",
                    description="Шаги Core Loop не ссылаются на описанные механики",
                    suggestion="Добавьте в описание шагов Core Loop ссылки на соответствующие механики",
                ))
        elif core_loop_present and not mechanics_present:
            issues.append(ConsistencyIssue(
                severity="info",
                section_a="core_loop",
                section_b="mechanics",
                issue_type="data_gap",
                description="Core Loop заполнен, но секция Mechanics отсутствует",
                suggestion="Заполните секцию Mechanics для согласованности с Core Loop",
            ))
        elif not core_loop_present and mechanics_present:
            issues.append(ConsistencyIssue(
                severity="error",
                section_a="core_loop",
                section_b="mechanics",
                issue_type="coreloop_mechanics_mismatch",
                description="Mechanics описаны, но Core Loop отсутствует — критическая секция",
                suggestion="Заполните секцию Core Loop, она является основой геймплея",
            ))

        # ── Пара 2: Progression ↔ Balance ──
        pair_key = "progression::balance"
        checked_pairs.append(pair_key)

        progression_present = _has_content("progression")
        balance_present = _has_content("balance")

        if progression_present and balance_present:
            prog_content = _get_content("progression").lower()
            bal_content = _get_content("balance").lower()

            # Проверяем перекрёстные ссылки на ключевые термины
            prog_has_levels = any(
                kw in prog_content for kw in ["уровень", "level", "тир", "tier", "ранг", "rank"]
            )
            bal_has_cost_power = any(
                kw in bal_content for kw in ["cost", "power", "стоимость", "мощность", "баланс", "balance"]
            )

            if not prog_has_levels:
                issues.append(ConsistencyIssue(
                    severity="warning",
                    section_a="progression",
                    section_b="balance",
                    issue_type="progression_balance_mismatch",
                    description="Прогрессия не описывает уровни/тиры, необходимые для балансировки",
                    suggestion="Добавьте описание уровней и тиров в секцию Progression",
                ))

            if not bal_has_cost_power:
                issues.append(ConsistencyIssue(
                    severity="info",
                    section_a="progression",
                    section_b="balance",
                    issue_type="progression_balance_mismatch",
                    description="Балансировка не содержит метрик cost/power",
                    suggestion="Добавьте таблицу cost-power в секцию Balance",
                ))
        elif progression_present and not balance_present:
            issues.append(ConsistencyIssue(
                severity="info",
                section_a="progression",
                section_b="balance",
                issue_type="data_gap",
                description="Прогрессия описана, но секция Balance отсутствует",
                suggestion="Заполните секцию Balance для проверки согласованности",
            ))
        elif not progression_present and balance_present:
            issues.append(ConsistencyIssue(
                severity="warning",
                section_a="progression",
                section_b="balance",
                issue_type="progression_balance_mismatch",
                description="Балансировка описана, но Progression отсутствует",
                suggestion="Заполните секцию Progression для согласования с балансом",
            ))

        # ── Пара 3: Economy ↔ Monetization ──
        pair_key = "economy::monetization"
        checked_pairs.append(pair_key)

        economy_present = _has_content("economy")
        monetization_present = _has_content("monetization")

        if economy_present and monetization_present:
            econ_content = _get_content("economy").lower()
            mon_content = _get_content("monetization").lower()

            # Проверяем: упоминаются ли ресурсы из экономики в монетизации
            econ_has_resources = any(
                kw in econ_content for kw in ["ресурс", "resource", "валют", "currency", "кристалл", "золото"]
            )
            mon_uses_resources = any(
                kw in mon_content for kw in ["ресурс", "resource", "валют", "currency", "кристалл", "золото", "покуп", "purchase", "premium"]
            )

            if econ_has_resources and not mon_uses_resources:
                issues.append(ConsistencyIssue(
                    severity="warning",
                    section_a="economy",
                    section_b="monetization",
                    issue_type="economy_monetization_mismatch",
                    description="Экономика описывает ресурсы, но монетизация не связана с ними",
                    suggestion="Укажите, какие ресурсы являются премиум-валютой в секции Monetization",
                ))
        elif economy_present and not monetization_present:
            issues.append(ConsistencyIssue(
                severity="info",
                section_a="economy",
                section_b="monetization",
                issue_type="data_gap",
                description="Экономика описана, но секция Monetization отсутствует",
                suggestion="Заполните секцию Monetization для полного описания бизнес-модели",
            ))
        elif not economy_present and monetization_present:
            issues.append(ConsistencyIssue(
                severity="warning",
                section_a="economy",
                section_b="monetization",
                issue_type="economy_monetization_mismatch",
                description="Монетизация описана без экономической модели",
                suggestion="Заполните секцию Economy для поддержки модели монетизации",
            ))

        # ── Пара 4: Narrative ↔ Mechanics ──
        pair_key = "narrative::mechanics"
        checked_pairs.append(pair_key)

        story_present = _has_content("story") or _has_content("lore")
        mechanics_present = _has_content("mechanics")

        if story_present and mechanics_present:
            story_content = (
                _get_content("story") + " " + _get_content("lore")
            ).lower()
            mech_content = _get_content("mechanics").lower()

            # Проверяем: противоречит ли тон истории механикам
            violent_mechanics = any(
                kw in mech_content for kw in ["бо", "combat", "атак", "attack", "урон", "damage", "стрельб", "shooting"]
            )
            peaceful_story = any(
                kw in story_content for kw in ["мир", "peace", "гармон", "harmony", "дружб", "friendship", "спокойн", "calm"]
            )

            if violent_mechanics and peaceful_story:
                issues.append(ConsistencyIssue(
                    severity="warning",
                    section_a="story",
                    section_b="mechanics",
                    issue_type="narrative_mechanics_mismatch",
                    description="Нарратив описывает мирный мир, но механики ориентированы на бой",
                    suggestion="Согласуйте тон истории с геймплейными механиками или добавьте мирные механики",
                ))
        elif story_present and not mechanics_present:
            issues.append(ConsistencyIssue(
                severity="info",
                section_a="story",
                section_b="mechanics",
                issue_type="data_gap",
                description="Нарратив описан, но механики отсутствуют",
                suggestion="Заполните секцию Mechanics для согласования с историей",
            ))

        # ── Пара 5: Лудонарративный диссонанс ──
        pair_key = "ludonarrative_dissonance"
        checked_pairs.append(pair_key)

        # Проверяем эстетику MDA и геймплей
        story_content = (_get_content("story") + " " + _get_content("lore")).lower()
        gameplay_content = (
            _get_content("core_loop") + " " + _get_content("mechanics") + " " + _get_content("progression")
        ).lower()

        if story_content and gameplay_content:
            # Определяем нарративный тон
            narrative_dark = any(
                kw in story_content for kw in ["мрачн", "dark", "ужас", "horror", "смерт", "death", "трагед", "tragedy", "отчаян", "despair"]
            )
            narrative_hopeful = any(
                kw in story_content for kw in ["надежд", "hope", "спасени", "salvation", "геро", "hero", "побед", "victory"]
            )
            gameplay_grindy = any(
                kw in gameplay_content for kw in ["повтор", "grind", "фарм", "farm", "монотон", "repetit"]
            )
            gameplay_casual = any(
                kw in gameplay_content for kw in ["казуал", "casual", "расслабл", "relax", "прост", "simple", "лёгк"]
            )

            if narrative_dark and gameplay_casual:
                issues.append(ConsistencyIssue(
                    severity="warning",
                    section_a="story",
                    section_b="core_loop",
                    issue_type="ludonarrative_dissonance",
                    description="Лудонарративный диссонанс: мрачный нарратив контрастирует с казуальным геймплеем",
                    suggestion="Измените тон нарратива или усложните механики для соответствия",
                ))
            elif narrative_hopeful and gameplay_grindy:
                issues.append(ConsistencyIssue(
                    severity="info",
                    section_a="story",
                    section_b="core_loop",
                    issue_type="ludonarrative_dissonance",
                    description="Потенциальный лудонарративный диссонанс: оптимистичный нарратив при гриндовом геймплеe",
                    suggestion="Убедитесь, что гринд обоснован историей (напр. тренировка героя)",
                ))

        # Подсчёт
        error_count = sum(1 for i in issues if i.severity == "error")
        warning_count = sum(1 for i in issues if i.severity == "warning")
        info_count = sum(1 for i in issues if i.severity == "info")

        return ConsistencyReport(
            issues=issues,
            error_count=error_count,
            warning_count=warning_count,
            info_count=info_count,
            is_valid=error_count == 0,
            checked_pairs=checked_pairs,
        )

    def _reconstruct_input_from_profile(self, profile: GDDProfile) -> GDDGenerationInput:
        """Восстановление GDDGenerationInput из GDDProfile для валидации."""
        return GDDGenerationInput(
            concept=None,
            core_loop=None,
            mda_profile=None,
            balance_result=None,
            progression_profile=None,
            economy_profile=None,
            target_format=profile.format_spec.format,
            target_audience_doc=profile.format_spec.audience,
            detail_level=profile.format_spec.detail_level,
        )

    # ========================================================
    # Этап 7: Форматирование документа (3.7.9)
    # ========================================================

    def format_document(
        self,
        assembled: GDDAssembledDocument,
        format_spec: GDDFormatSpec,
        input_data: GDDGenerationInput,
    ) -> GDDFormattedDocument:
        """
        Этап 7: Форматирование собранного GDD в Markdown.

        Алгоритм 3.7.9:
        1. Заголовок документа
        2. Подзаголовок с мета-информацией
        3. Оглавление
        4. Секции с нумерацией
        5. Подсчёт слов и страниц
        """
        start = time.time()

        # Заголовок
        title = "Game Design Document"
        if input_data.concept:
            title = input_data.concept.get("title", title) or title

        # Дата
        date_str = datetime.now().strftime("%Y-%m-%d")

        # Жанр
        genre = ""
        if input_data.concept:
            genre = input_data.concept.get("genre", "")

        # Формат документа
        format_name = format_spec.format.replace("_", " ").title()

        # ── Строим Markdown ──
        md_lines: list[str] = []

        # Заголовок
        md_lines.append(f"# {title}")
        md_lines.append("")

        # Подзаголовок
        subtitle_parts = [format_name]
        if genre:
            subtitle_parts.append(genre)
        subtitle_parts.append(date_str)
        md_lines.append(f"*{' | '.join(subtitle_parts)}*")
        md_lines.append("")
        md_lines.append("---")
        md_lines.append("")

        # Оглавление
        toc_lines: list[str] = ["## Оглавление", ""]
        section_counter = 0

        for section_name in assembled.section_order:
            section = assembled.sections.get(section_name)
            if not section:
                continue

            section_counter += 1
            display_name = self._section_display_name(section_name)
            anchor = section_name.replace("_", "-").lower()
            toc_lines.append(f"{section_counter}. [{display_name}](#{anchor})")

        toc_lines.append("")

        md_lines.extend(toc_lines)
        md_lines.append("---")
        md_lines.append("")

        # Секции
        section_num = 0
        for section_name in assembled.section_order:
            section = assembled.sections.get(section_name)
            if not section:
                continue

            section_num += 1
            display_name = self._section_display_name(section_name)

            md_lines.append(f"## {section_num}. {display_name}")
            md_lines.append("")

            if section.content:
                md_lines.append(section.content)
            else:
                md_lines.append("*Раздел не заполнен*")

            md_lines.append("")

            # Добавляем диаграмму если есть
            # (диаграмма уже встроена в content через SectionContent.diagram)

            # Отметка источника
            if section.source and section.source != "auto_fill":
                source_labels = {
                    "ai_enrich": "🤖 AI-обогащено",
                    "ai_generate": "🤖 AI-сгенерировано",
                    "manual": "✏️ Ручное заполнение",
                }
                label = source_labels.get(section.source, "")
                if label:
                    md_lines.append(f"*{label}*")
                    md_lines.append("")

        # Подвал
        md_lines.append("---")
        md_lines.append(f"*Документ сгенерирован: {date_str} | Формат: {format_name}*")

        markdown = "\n".join(md_lines)

        # Подсчёт слов и страниц
        word_count = len(markdown.split())
        estimated_pages = max(1, word_count // 250)

        # Оглавление как строка
        table_of_contents = "\n".join(toc_lines)

        result = GDDFormattedDocument(
            markdown=markdown,
            title=title,
            table_of_contents=table_of_contents,
            section_count=section_num,
            word_count=word_count,
            estimated_pages=estimated_pages,
        )

        logger.info(
            f"[Stage 7] Formatted: {section_num} sections, "
            f"{word_count} words, ~{estimated_pages} pages "
            f"({time.time() - start:.2f}s)"
        )

        return result

    def _section_display_name(self, section_name: str) -> str:
        """Получить человекочитаемое название секции."""
        name_map = {
            "title": "Название",
            "logline": "Логлайн",
            "overview": "Обзор",
            "genre_platform": "Жанр и платформы",
            "target_audience": "Целевая аудитория",
            "uniqueness": "Уникальность",
            "license": "Лицензия",
            "core_loop": "Core Loop",
            "controls": "Управление",
            "mechanics": "Механики",
            "camera_perspective": "Камера",
            "progression": "Прогрессия",
            "balance": "Балансировка",
            "difficulty": "Сложность",
            "game_modes": "Игровые режимы",
            "characters": "Персонажи",
            "story": "Сюжет",
            "dialogues": "Диалоги",
            "quests": "Квесты",
            "lore": "Лор",
            "world_structure": "Структура мира",
            "level_design": "Дизайн уровней",
            "navigation": "Навигация",
            "combat_spaces": "Боевые пространства",
            "resources": "Ресурсы",
            "economy": "Экономика",
            "tech_tree": "Дерево технологий",
            "difficulty_curve": "Кривая сложности",
            "hud_ui": "HUD и UI",
            "menus": "Меню",
            "visual_style": "Визуальный стиль",
            "sound_music": "Звук и музыка",
            "multiplayer_modes": "Мультиплеер",
            "social_features": "Социальные функции",
            "meta_game": "Мета-игра",
            "tech_requirements": "Технические требования",
            "platform_ports": "Платформы",
            "monetization": "Монетизация",
            "milestones_budget": "Майлстоуны и бюджет",
            "title_logline": "Название и логлайн",
            "game_type": "Тип игры",
            "originality": "Оригинальность",
            "feasibility": "Осуществимость",
            "visual_hook": "Визуальный крючок",
            "content_overview": "Обзор контента",
            "player_experience_goal": "Целевой опыт игрока",
            "system_map": "Карта систем",
            "feedback_patterns": "Паттерны обратной связи",
            "success_metrics": "Метрики успеха",
            "storyline_map": "Карта сюжета",
            "cutscene_scripts": "Сценарии кат-сцен",
            "quest_matrix": "Матрица квестов",
            "lore_db": "База лора",
            "dissension_validator": "Валидация диссонанса",
            "pitch_deck": "Pitch Deck",
            "risks": "Риски",
            "team": "Команда",
            "concept_overview": "Обзор концепции",
            "mda_analysis": "MDA-анализ",
            "balance_tables": "Таблицы баланса",
            "progression_curves": "Кривые прогрессии",
            "economy_model": "Модель экономики",
            "character_profiles": "Профили персонажей",
            "narrative_arc": "Нарративная арка",
            "ui_wireframes": "Вайрфреймы UI",
            "technical_specs": "Технические спецификации",
            "checklist_results": "Результаты чек-листов",
        }
        return name_map.get(section_name, section_name.replace("_", " ").title())

    # ========================================================
    # Этап 8: Экспорт GDD
    # ========================================================

    async def export_gdd(
        self,
        formatted: GDDFormattedDocument,
        export_format: ExportFormat,
        project_title: str = "GDD",
    ) -> GDDExportResult:
        """
        Этап 8: Экспорт GDD в выбранный формат.

        Поддерживаемые форматы:
        - md: Markdown (контент напрямую)
        - html: HTML с CSS-стилями
        - pdf: PDF через WeasyPrint (fallback → HTML)
        - docx: DOCX через python-docx
        """
        start = time.time()

        # Безопасное имя файла
        safe_title = re.sub(r'[^\w\s-]', '', project_title).strip().replace(' ', '_')[:50]

        try:
            if export_format == "md":
                result = self._export_markdown(formatted, safe_title)

            elif export_format == "html":
                result = self._export_html(formatted, safe_title)

            elif export_format == "pdf":
                result = self._export_pdf(formatted, safe_title)

            elif export_format == "docx":
                result = self._export_docx(formatted, safe_title)

            else:
                result = GDDExportResult(
                    format=export_format,
                    success=False,
                    error_message=f"Неподдерживаемый формат экспорта: {export_format}",
                )

        except Exception as e:
            logger.error(f"[Stage 8] Export failed for {export_format}: {e}", exc_info=True)
            result = GDDExportResult(
                format=export_format,
                success=False,
                error_message=f"Ошибка экспорта: {str(e)}",
            )

        latency = time.time() - start
        logger.info(
            f"[Stage 8] Export to {export_format}: "
            f"success={result.success}, "
            f"size={result.size_bytes} bytes "
            f"({latency:.2f}s)"
        )

        return result

    def _export_markdown(
        self, formatted: GDDFormattedDocument, safe_title: str
    ) -> GDDExportResult:
        """Экспорт в Markdown."""
        content = formatted.markdown
        return GDDExportResult(
            format="md",
            content=content,
            file_name=f"{safe_title}.md",
            content_type="text/markdown",
            size_bytes=len(content.encode("utf-8")),
            success=True,
        )

    def _export_html(
        self, formatted: GDDFormattedDocument, safe_title: str
    ) -> GDDExportResult:
        """Экспорт в HTML с профессиональным CSS-оформлением."""
        html_content = self._markdown_to_html(formatted.markdown, formatted.title)
        return GDDExportResult(
            format="html",
            content=html_content,
            file_name=f"{safe_title}.html",
            content_type="text/html",
            size_bytes=len(html_content.encode("utf-8")),
            success=True,
        )

    def _export_pdf(
        self, formatted: GDDFormattedDocument, safe_title: str
    ) -> GDDExportResult:
        """Экспорт в PDF через WeasyPrint (fallback → HTML)."""
        html_content = self._markdown_to_html(formatted.markdown, formatted.title)

        try:
            from weasyprint import HTML as WeasyHTML  # type: ignore

            tmp_dir = tempfile.mkdtemp(prefix="gdd_export_")
            file_path = os.path.join(tmp_dir, f"{safe_title}.pdf")

            WeasyHTML(string=html_content).write_pdf(file_path)

            size_bytes = os.path.getsize(file_path) if os.path.exists(file_path) else 0

            return GDDExportResult(
                format="pdf",
                content="",
                file_path=file_path,
                file_name=f"{safe_title}.pdf",
                content_type="application/pdf",
                size_bytes=size_bytes,
                success=True,
            )

        except ImportError:
            logger.warning("[Stage 8] WeasyPrint not installed, falling back to HTML export")
            return GDDExportResult(
                format="pdf",
                content=html_content,
                file_name=f"{safe_title}.html",
                content_type="text/html",
                size_bytes=len(html_content.encode("utf-8")),
                success=True,
                error_message="WeasyPrint не установлен. Экспортировано как HTML.",
            )

    def _export_docx(
        self, formatted: GDDFormattedDocument, safe_title: str
    ) -> GDDExportResult:
        """Экспорт в DOCX через python-docx."""
        try:
            from docx import Document  # type: ignore
            from docx.shared import Pt, Inches  # type: ignore

            doc = Document()

            # Заголовок
            doc.add_heading(formatted.title, level=0)

            # Секции
            lines = formatted.markdown.split("\n")
            current_section = ""

            for line in lines:
                stripped = line.strip()

                if not stripped:
                    continue

                # Заголовки
                if stripped.startswith("## "):
                    heading_text = stripped[3:].strip()
                    # Убираем якорные ссылки
                    heading_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', heading_text)
                    doc.add_heading(heading_text, level=2)
                    current_section = heading_text

                elif stripped.startswith("# "):
                    # Основной заголовок — пропускаем, уже добавлен
                    pass

                elif stripped.startswith("---"):
                    doc.add_paragraph("")  # Разделитель

                elif stripped.startswith("- ") or stripped.startswith("* "):
                    # Список
                    item_text = stripped[2:].strip()
                    item_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', item_text)
                    doc.add_paragraph(item_text, style="List Bullet")

                elif re.match(r'^\d+\.\s', stripped):
                    # Нумерованный список
                    item_text = re.sub(r'^\d+\.\s', '', stripped)
                    item_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', item_text)
                    doc.add_paragraph(item_text, style="List Number")

                elif stripped.startswith("```"):
                    # Блок кода — пропускаем для DOCX
                    pass

                elif stripped.startswith("*") and stripped.endswith("*") and len(stripped) > 2:
                    # Курсивный текст
                    text = stripped.strip("*")
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.italic = True

                else:
                    # Обычный текст
                    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', stripped)
                    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
                    if text and text != "---":
                        doc.add_paragraph(text)

            # Сохраняем
            tmp_dir = tempfile.mkdtemp(prefix="gdd_export_")
            file_path = os.path.join(tmp_dir, f"{safe_title}.docx")
            doc.save(file_path)

            size_bytes = os.path.getsize(file_path) if os.path.exists(file_path) else 0

            return GDDExportResult(
                format="docx",
                content="",
                file_path=file_path,
                file_name=f"{safe_title}.docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                size_bytes=size_bytes,
                success=True,
            )

        except ImportError:
            logger.warning("[Stage 8] python-docx not installed, returning error")
            return GDDExportResult(
                format="docx",
                success=False,
                error_message="python-docx не установлен. Установите: pip install python-docx",
            )

    def _markdown_to_html(self, markdown: str, title: str) -> str:
        """Простая конвертация Markdown → HTML с профессиональным CSS."""
        html = markdown

        # Блоки кода (mermaid и обычные)
        html = re.sub(
            r'```(\w*)\n(.*?)```',
            r'<pre><code class="\1">\2</code></pre>',
            html,
            flags=re.DOTALL,
        )

        # Заголовки
        html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)

        # Жирный и курсив
        html = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', html)
        html = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', html)

        # Ссылки
        html = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', html)

        # Нумерованные списки
        html = re.sub(
            r'^(\d+)\.\s+(.+)$',
            r'<li>\2</li>',
            html,
            flags=re.MULTILINE,
        )

        # Маркированные списки
        html = re.sub(r'^- (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)

        # Оборачиваем li в ul/ol (упрощённо)
        html = re.sub(
            r'(<li>.*?</li>(\n<li>.*?</li>)*)',
            r'<ul>\1</ul>',
            html,
            flags=re.DOTALL,
        )

        # Горизонтальные линии
        html = re.sub(r'^---+$', r'<hr>', html, flags=re.MULTILINE)

        # Параграфы (двойной перенос строки)
        html = re.sub(r'\n\n+', r'\n</p>\n<p>\n', html)

        # Таблицы (упрощённая конвертация)
        html = self._convert_markdown_tables_to_html(html)

        css = """
        <style>
            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                max-width: 900px;
                margin: 0 auto;
                padding: 40px 60px;
                color: #1a1a1a;
                line-height: 1.7;
                background: #ffffff;
            }
            h1 {
                font-size: 28pt;
                color: #1a3a5c;
                border-bottom: 3px solid #1a3a5c;
                padding-bottom: 12px;
                margin-top: 0;
            }
            h2 {
                font-size: 18pt;
                color: #2c5f8a;
                border-bottom: 1px solid #e0e0e0;
                padding-bottom: 6px;
                margin-top: 36px;
            }
            h3 {
                font-size: 14pt;
                color: #3a7cb8;
                margin-top: 24px;
            }
            p {
                margin: 8px 0;
                text-align: justify;
            }
            strong { color: #1a3a5c; }
            em { color: #555; }
            ul, ol { margin: 8px 0; padding-left: 24px; }
            li { margin: 4px 0; }
            table {
                border-collapse: collapse;
                width: 100%;
                margin: 16px 0;
                font-size: 10pt;
            }
            th, td {
                border: 1px solid #ccc;
                padding: 8px 12px;
                text-align: left;
            }
            th {
                background: #2c5f8a;
                color: white;
                font-weight: 600;
            }
            tr:nth-child(even) { background: #f5f8fc; }
            pre {
                background: #f4f4f4;
                padding: 16px;
                border-radius: 4px;
                overflow-x: auto;
                font-size: 9pt;
            }
            code { font-family: 'Consolas', 'Monaco', monospace; }
            hr {
                border: none;
                border-top: 1px solid #e0e0e0;
                margin: 24px 0;
            }
            a { color: #2c5f8a; text-decoration: none; }
            @media print {
                body { max-width: none; padding: 20px; }
                h2 { page-break-before: auto; }
            }
        </style>
        """

        return f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    {css}
</head>
<body>
<p>
{html}
</p>
</body>
</html>"""

    def _convert_markdown_tables_to_html(self, html: str) -> str:
        """Конвертация Markdown-таблиц в HTML."""
        # Ищем блоки таблиц (строки с |)
        table_pattern = re.compile(
            r'((?:^\|.*\|$)\n?)+',
            re.MULTILINE,
        )

        def _table_replacer(match: re.Match) -> str:
            table_text = match.group(0)
            rows = [
                line.strip()
                for line in table_text.strip().split("\n")
                if line.strip() and line.strip().startswith("|")
            ]

            if len(rows) < 2:
                return table_text

            html_rows: list[str] = []

            for i, row in enumerate(rows):
                cells = [c.strip() for c in row.split("|")[1:-1]]

                # Пропускаем разделитель (|---|---|)
                if all(re.match(r'^[-:]+$', c) for c in cells):
                    continue

                tag = "th" if i == 0 else "td"
                cell_html = "".join(
                    f"<{tag}>{c}</{tag}>" for c in cells
                )
                html_rows.append(f"<tr>{cell_html}</tr>")

            if not html_rows:
                return table_text

            return (
                "<table>\n<thead>\n"
                + html_rows[0]
                + "\n</thead>\n<tbody>\n"
                + "\n".join(html_rows[1:])
                + "\n</tbody>\n</table>"
            )

        return table_pattern.sub(_table_replacer, html)

    # ========================================================
    # Полный пайплайн: Этапы 1–8
    # ========================================================

    async def generate_stages_1_8(
        self,
        input_data: GDDGenerationInput,
    ) -> GDDProfile:
        """
        Полный пайплайн генерации GDD — Этапы 1–7 алгоритма 3.7.

        Выполняет:
        1–5: Стандартный пайплайн (format → map → auto_fill → ai → manual)
        6: Сшивка и валидация документа
        7: Форматирование в Markdown

        Этап 8 (экспорт) вызывается отдельно через export_gdd().

        Returns:
            GDDProfile с результатами всех семи этапов
        """
        pipeline_start = time.time()

        # === Этапы 1–5 ===
        profile = await self.generate_stages_1_5(input_data)

        # === Этап 6: Сшивка и валидация ===
        assembled = await self.assemble_gdd(profile)
        profile.assembled_document = assembled

        # === Этап 7: Форматирование ===
        formatted = self.format_document(assembled, profile.format_spec, input_data)
        profile.formatted_document = formatted

        # Обновляем метаданные профиля
        profile.stages_completed = [1, 2, 3, 4, 5, 6, 7]
        profile.coverage_score = assembled.coverage_score

        latency_ms = int((time.time() - pipeline_start) * 1000)
        profile.latency_ms = latency_ms

        logger.info(
            f"[Pipeline 1-8] Completed in {latency_ms}ms. "
            f"Format: {profile.format_spec.format}, "
            f"Sections: {assembled.total_sections}, "
            f"Filled: {assembled.filled_sections}, "
            f"Coverage: {assembled.coverage_score:.2f}, "
            f"Consistency: {assembled.consistency_report.error_count} errors, "
            f"Words: {formatted.word_count}, "
            f"Pages: ~{formatted.estimated_pages}"
        )

        return profile
